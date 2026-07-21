# ==========================================================
# ALBAT RAPPORT — Génération de rapport type (étude d'impact)
# ==========================================================
#
# Ce module s'inspire de la trame "Silva Environnement" (suivi
# acoustique en nacelle) et se construit par étapes :
#
#   1. Profil du bureau d'étude (logo + coordonnées), enregistrable
#      et rechargeable — c'est cette étape qui est implémentée ici.
#   2. Génération de la page de garde à partir du profil.
#   3. Injection automatique des résultats des autres modules
#      (Correlations / Graph / Bridage / Scenar).
#   4. Zones de texte à compléter par l'expert (interprétation).
#
# ==========================================================

import sys
import os
import re
import json

try:
    import openpyxl
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False
import shutil
import subprocess
import datetime as dtmod
from io import BytesIO

try:
    from PIL import Image as PILImage
    PIL_OK = True
except ImportError:
    PIL_OK = False


def open_file(path):
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.call(["open", path])
        else:
            subprocess.call(["xdg-open", path])
    except Exception:
        pass

try:
    from docx import Document
    from docx.document import Document as _DocxDocument
    from docx.shared import Cm, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


def _iter_block_items(parent):
    """Itère les paragraphes ET les tableaux d'un document (ou
    d'une cellule) source dans leur ordre d'apparition réel.
    Nécessaire pour ne pas perdre les tableaux natifs (ex : grille
    de sensibilité SFEPM) lors de la copie d'un fichier .docx
    fourni par l'utilisateur (matériel utilisé, statuts des
    espèces) : .paragraphs seul les ignore complètement."""

    if isinstance(parent, _DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("parent non supporté pour _iter_block_items")

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)

from PySide6.QtWidgets import (
    QApplication, QWidget, QLabel, QPushButton,
    QFileDialog, QVBoxLayout, QHBoxLayout, QMessageBox,
    QLineEdit, QTextEdit, QScrollArea, QFrame, QFontComboBox,
    QSpinBox, QColorDialog, QComboBox, QListWidget, QListWidgetItem,
    QCheckBox
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QPixmap, QColor, QFont

from modules.utils import resource_path

ASSETS = resource_path("assets")
BG = resource_path("assets", "background_graph.png")

BIBLIO_DIR = resource_path("dossiers pour rapport", "bibliographie")
MATERIEL_DIR = resource_path("dossiers pour rapport", "materiel utilisé")
STATUTS_DIR = resource_path("dossiers pour rapport", "Statuts des espèces")
ANALYSE_DIR = resource_path(
    "dossiers pour rapport", "Analyse des enregistrements"
)
IMAGES_DIR = resource_path("dossiers pour rapport", "images")


# ==============================================================
# TEXTES GÉNÉRIQUES SPÉCIFIQUES AU BUREAU "SILVA ENVIRONNEMENT"
# ==============================================================
# Rassemblés ici, à part du code de génération, pour que ce qui
# relève du style rédactionnel propre à Silva Environnement (et
# non d'une logique générique réutilisable par n'importe quel
# bureau) soit facile à repérer et à extraire le jour où un
# deuxième bureau, avec une structure de rapport différente, doit
# être pris en charge. Ce bloc est un simple rangement : il ne
# change rien au comportement actuel du générateur.
#
# Convention : un texte contenant des parties variables (nom du
# parc, année...) est un gabarit à utiliser avec .format(model=...) ;
# les paramètres attendus sont indiqués entre {accolades}.
# ==============================================================

TEXTES_SILVA = {

    # --- Introduction ---
    "introduction_p1": (
        "Situé en région [RÉGION], le parc éolien de {parc} se "
        "compose de [NOMBRE] éoliennes [MODÈLE] d'une capacité de "
        "[PUISSANCE] MW chacune. Dans le cadre de [CONTEXTE DU "
        "SUIVI — mise en service / repowering / suivi "
        "réglementaire...], [MAÎTRE D'OUVRAGE] a fait réaliser un "
        "suivi environnemental en {annee}."
    ),
    "introduction_p2": (
        "Ce document présente les résultats issus du suivi "
        "acoustique en nacelle mené en {annee} sur le parc éolien "
        "de {parc}."
    ),
    "introduction_todo": (
        "[Texte générique à adapter — complétez ou reformulez ce "
        "paragraphe selon le contexte réel du projet (région, "
        "nombre et modèle d'éoliennes, puissance, maître "
        "d'ouvrage, raison du suivi).]"
    ),

    # --- I. Matériel et méthode (introduction générale) ---
    "materiel_methode_p1": (
        "Le suivi automatique en altitude apparait comme une "
        "formidable avancée technologique dans la perception de "
        "l'activité des chauves-souris en continu."
    ),
    "materiel_methode_p2": (
        "Il se justifie d'abord par la grande disparité d'activité "
        "altitudinale (phénomène d'autant plus marqué en milieu "
        "boisé entre une situation en sous-bois ou au-dessus de la "
        "canopée)."
    ),
    "materiel_methode_p3": (
        "Il permet également de rechercher efficacement "
        "l'éventuelle présence d'une activité migratoire, de "
        "transit ou bien de haut vol, perception très difficile "
        "depuis le sol selon les espèces et selon les obstacles "
        "acoustiques (canopée par exemple)."
    ),
    "materiel_methode_p4": (
        "Enfin, il représente une réponse adaptée aux importants "
        "biais de l'échantillonnage ponctuel quand on sait combien "
        "l'activité des chauves-souris est très hétérogène dans le "
        "temps (d'une nuit à l'autre) sous l'influence d'un cumul "
        "de facteurs bioclimatiques."
    ),

    # --- I.1. Matériel utilisé (introduction générique) ---
    "materiel_utilise_intro": (
        "Le suivi de l'activité chiroptérologique en altitude est "
        "réalisé à l'aide d'un enregistreur ultrasonore installé "
        "dans la nacelle de l'éolienne, permettant un "
        "fonctionnement autonome sur l'ensemble de la période "
        "d'étude et une détection représentative de l'activité "
        "réellement présente au niveau du rotor."
    ),

    # --- II.4. Influence de la température ---
    "influence_temperature_p1": (
        "Concernant la température, l'activité est globalement "
        "plus marquée à partir de 16°C (Loiret Nature "
        "Environnement, 2009) avec une augmentation de l'activité "
        "et des valeurs comprises entre 10 à 25°C (Brinkmann, "
        "2011). Ces données sont néanmoins dépendantes de la "
        "situation géographique et de l'altitude."
    ),
    "influence_temperature_p2": (
        "A noter que la tolérance à la température est variable "
        "selon les espèces (Ecosphère, 2017). La Pipistrelle de "
        "Nathusius et la Pipistrelle commune semblent ainsi encore "
        "mobiles lors de faibles températures. Leur plus basse "
        "activité a été mesurée respectivement à 2°C et 1°C "
        "(Joiris, 2012)."
    ),

    # --- II.4. Influence de la vitesse de vent ---
    "influence_vent_p1": (
        "La vitesse du vent apparaît comme un facteur clé de "
        "régulation de l'activité des chauves-souris en altitude. "
        "Des études ont par exemple montré que 94% des contacts "
        "sont enregistrés pour des vitesses de vent inférieures à "
        "6 m/s (Loiret Nature Environnement, 2009) ou 6,5 m/s "
        "(Behr, 2007). Ainsi, l'activité est plus élevée en "
        "période de faible vent."
    ),

    # --- II.4. Influence de l'heure de la nuit ---
    "influence_heure_nuit_p1": (
        "Le risque de collision des chauves-souris avec les "
        "éoliennes varie en fonction de leur activité, qu'elle "
        "soit liée aux périodes de chasse ou de transit. Ainsi, "
        "plus les chauves-souris sont actives, plus le risque de "
        "collision augmente. Différentes études quantifient "
        "l'importance du début de la nuit (les 3 premières heures "
        "en général). D'autres études ont mis l'accent sur le "
        "premier quart de la nuit (Brinkmann, 2011) voire le "
        "premier tiers de la nuit (Behr, 2007)."
    ),
    "influence_heure_nuit_p2": (
        "Une majorité d'espèces montre une phénologie horaire "
        "marquée avec un net pic d'activité dans les 2 premières "
        "heures de la nuit (Haquart, 2012). L'activité baisse "
        "ensuite de manière plus ou moins constante (Brinkmann, "
        "2011) et serait ainsi plus faible vers la fin de la nuit, "
        "c'est-à-dire 4h à 7h après le coucher du soleil "
        "(Marchais, 2010)."
    ),
}


class RapportWindow(QWidget):

    def __init__(
        self,
        park_name="",
        ref_city="",
        ref_year=None,
        bureau_etude="",
        default_save_dir="",
        graph_dir_hint="",
        bureau_profile_path="",
        scenar_dir_hint="",
        qgis_dir_hint=""
    ):
        super().__init__()

        self.park_name = park_name
        self.ref_city = ref_city
        self.ref_year = ref_year
        self.bureau_etude = bureau_etude
        self.default_save_dir = default_save_dir
        self.bureau_profile_path = bureau_profile_path

        # Profil du bureau d'étude actuellement chargé
        self.logo_path = None
        self.logo_secondaire_path = None
        self.profile_path = None
        self.graphs_dir = (
            graph_dir_hint
            if graph_dir_hint and os.path.isdir(graph_dir_hint)
            else None
        )
        self.scenar_dir = (
            scenar_dir_hint
            if scenar_dir_hint and os.path.isdir(scenar_dir_hint)
            else None
        )
        self.qgis_dir = (
            qgis_dir_hint
            if qgis_dir_hint and os.path.isdir(qgis_dir_hint)
            else None
        )
        self.materiel_dir = None
        self.materiel_files = {}
        self.materiel_figure_files = {}
        self.biblio_file = None

        self.setWindowTitle("Albat Rapport")
        self.resize(700, 1000)
        self.setMinimumSize(600, 900)

        self.setStyleSheet("""
        QWidget{
            background:transparent;
            color:white;
            font-family:"Helvetica Neue","Segoe UI",Arial,sans-serif;
        }

        QLabel{
            font-size:14px;
        }

        QLineEdit, QTextEdit{
            background:black;
            color:white;
            border:1px solid rgba(255,255,255,60);
            border-radius:10px;
            padding:8px;
            font-size:13px;
        }

        QPushButton{
            background:rgba(90,127,71,145);
            border:1px solid rgba(255,255,255,40);
            border-radius:14px;
            padding:10px;
            font-size:14px;
            font-weight:700;
            color:white;
        }

        QPushButton:hover{
            background:rgba(130,184,99,220);
        }
        """)

        self.bg = QLabel(self)
        self.bg.setGeometry(self.rect())
        self.bg.setAlignment(Qt.AlignCenter)

        if os.path.exists(BG):
            p = QPixmap(BG)
            if not p.isNull():
                self.bg.setPixmap(
                    p.scaled(
                        self.size(),
                        Qt.KeepAspectRatioByExpanding,
                        Qt.SmoothTransformation
                    )
                )

        self.bg.lower()

        self.build_ui()

        # Lecture automatique des dossiers intégrés au démarrage
        # (pas besoin de cliquer sur "Rafraîchir" la première fois).
        self.pick_materiel_dir()
        self.pick_biblio_file()

        # Si un profil de bureau d'étude a été choisi au démarrage
        # du logiciel, on le précharge ici — load_profile_from_path
        # se charge elle-même de replier le bloc si des champs
        # essentiels sont remplis.
        if self.bureau_profile_path and os.path.exists(self.bureau_profile_path):
            try:
                self.load_profile_from_path(self.bureau_profile_path)
            except Exception:
                pass

    # ======================================================

    def resizeEvent(self, event):

        self.bg.setGeometry(self.rect())

        if os.path.exists(BG):
            p = QPixmap(BG)
            if not p.isNull():
                self.bg.setPixmap(
                    p.scaled(
                        self.size(),
                        Qt.KeepAspectRatioByExpanding,
                        Qt.SmoothTransformation
                    )
                )

        super().resizeEvent(event)

    # ======================================================

    def _make_collapsible_header(self, title_text):
        """
        Construit un en-tête de carte repliable : flèche de
        repli/déploi tout à gauche, titre juste à côté.
        Retourne (layout, bouton_flèche).
        """

        header_row = QHBoxLayout()
        header_row.setSpacing(8)

        btn_toggle = QPushButton("▼")
        btn_toggle.setFixedSize(28, 28)
        btn_toggle.setStyleSheet("""
            QPushButton{
                background:rgba(90,127,71,180);
                border:1px solid rgba(255,255,255,60);
                border-radius:14px;
                font-size:12px;
                font-weight:900;
                padding:0px;
            }
            QPushButton:hover{
                background:rgba(130,184,99,220);
            }
        """)

        card_title = QLabel(title_text)
        card_title.setAlignment(Qt.AlignCenter)
        card_title.setStyleSheet("""
            color: #E6E8C8;
            font-weight: 700;
            font-size: 13px;
            letter-spacing: 1px;
            background-color: rgba(0, 0, 0, 0.50);
            border: 1px solid rgba(140,170,90,0.60);
            border-radius: 12px;
            padding: 4px 20px;
        """)

        header_row.addWidget(btn_toggle)
        header_row.addWidget(card_title, 1)

        return header_row, btn_toggle

    # ======================================================

    def build_ui(self):

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea{background:transparent; border:none;}
            QScrollBar:vertical{
                background: rgba(20,20,20,220);
                width: 14px;
                margin: 0px;
                border-radius: 7px;
            }
            QScrollBar::handle:vertical{
                background: rgba(130,184,99,230);
                min-height: 30px;
                border-radius: 7px;
            }
            QScrollBar::handle:vertical:hover{
                background: rgba(150,204,110,255);
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical{
                height: 0px;
                border: none;
            }
            QScrollBar::add-page:vertical,
            QScrollBar::sub-page:vertical{
                background: none;
            }
        """)

        content = QWidget()
        content.setStyleSheet("background:transparent;")

        main = QVBoxLayout(content)
        main.setContentsMargins(20, 20, 20, 20)
        main.setSpacing(12)

        # Le grand titre décoratif "Albat / RAPPORT" a été retiré :
        # l'onglet actif dans la barre d'onglets du haut suffit
        # déjà à indiquer où l'on se trouve.

        sub = QLabel(
            "Profil du bureau d'étude — logo et coordonnées utilisés "
            "pour la page de garde des rapports générés."
        )
        sub.setWordWrap(True)
        sub.setAlignment(Qt.AlignCenter)
        sub.setStyleSheet(
            "background:transparent; color:#cfe0a0; font-size:11px;"
        )
        main.addWidget(sub)

        # ==========================================
        # CARTE : PROFIL DU BUREAU D'ÉTUDE
        # ==========================================

        profile_card = QFrame()
        profile_card.setStyleSheet("""
            QFrame{
                background:rgba(20,35,15,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        profile_outer = QVBoxLayout(profile_card)
        profile_outer.setContentsMargins(16, 16, 16, 16)
        profile_outer.setSpacing(10)

        # --- En-tête cliquable (flèche à gauche + titre) ---
        header_row, self.btn_toggle_profile = self._make_collapsible_header(
            "LOGO ET COORDONNÉES DU BUREAU"
        )
        self.btn_toggle_profile.clicked.connect(self.toggle_profile_section)

        profile_outer.addLayout(header_row)

        # --- Contenu repliable ---
        self.profile_content = QWidget()
        self.profile_content.setStyleSheet("background:transparent;")

        profile_layout = QVBoxLayout(self.profile_content)
        profile_layout.setContentsMargins(0, 0, 0, 0)
        profile_layout.setSpacing(10)

        # --- Logo principal (page de garde) ---
        lab_logo_principal = QLabel("Logo principal (page de garde)")
        lab_logo_principal.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        profile_layout.addWidget(lab_logo_principal)

        logo_row = QHBoxLayout()
        logo_row.setSpacing(12)

        self.logo_preview = QLabel("Aucun logo")
        self.logo_preview.setFixedSize(120, 120)
        self.logo_preview.setAlignment(Qt.AlignCenter)
        self.logo_preview.setStyleSheet("""
            background:rgba(0,0,0,0.35);
            border:1px solid rgba(255,255,255,60);
            border-radius:10px;
            color:#cccccc;
            font-size:11px;
        """)

        logo_btn_col = QVBoxLayout()
        logo_btn_col.setSpacing(8)

        btn_logo = QPushButton("Choisir un logo (image)")
        btn_logo.clicked.connect(self.pick_logo)

        self.lbl_logo_status = QLabel("Aucun fichier sélectionné")
        self.lbl_logo_status.setWordWrap(True)
        self.lbl_logo_status.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        lbl_logo_format = QLabel(
            "Format recommandé : environ 500×500px, fond transparent "
            "de préférence (.png)"
        )
        lbl_logo_format.setWordWrap(True)
        lbl_logo_format.setStyleSheet(
            "background:transparent; color:#777777; font-size:9px; "
            "font-style:italic;"
        )

        logo_btn_col.addWidget(btn_logo)
        logo_btn_col.addWidget(self.lbl_logo_status)
        logo_btn_col.addWidget(lbl_logo_format)
        logo_btn_col.addStretch()

        logo_row.addWidget(self.logo_preview)
        logo_row.addLayout(logo_btn_col, 1)

        profile_layout.addLayout(logo_row)

        # --- Logo secondaire (petit logo au-dessus des coordonnées) ---
        lab_logo_secondaire = QLabel(
            "Logo secondaire (petit logo au-dessus des coordonnées, "
            "en bas de la page de garde)"
        )
        lab_logo_secondaire.setWordWrap(True)
        lab_logo_secondaire.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        profile_layout.addWidget(lab_logo_secondaire)

        logo2_row = QHBoxLayout()
        logo2_row.setSpacing(12)

        self.logo2_preview = QLabel("Aucun logo")
        self.logo2_preview.setFixedSize(120, 120)
        self.logo2_preview.setAlignment(Qt.AlignCenter)
        self.logo2_preview.setStyleSheet("""
            background:rgba(0,0,0,0.35);
            border:1px solid rgba(255,255,255,60);
            border-radius:10px;
            color:#cccccc;
            font-size:11px;
        """)

        logo2_btn_col = QVBoxLayout()
        logo2_btn_col.setSpacing(8)

        btn_logo2 = QPushButton("Choisir un logo secondaire (image)")
        btn_logo2.clicked.connect(self.pick_logo_secondaire)

        self.lbl_logo2_status = QLabel(
            "Aucun fichier sélectionné (le logo principal sera "
            "réutilisé en miniature par défaut)"
        )
        self.lbl_logo2_status.setWordWrap(True)
        self.lbl_logo2_status.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        lbl_logo2_format = QLabel(
            "Format recommandé : environ 150×150px, fond transparent "
            "de préférence (.png)"
        )
        lbl_logo2_format.setWordWrap(True)
        lbl_logo2_format.setStyleSheet(
            "background:transparent; color:#777777; font-size:9px; "
            "font-style:italic;"
        )

        logo2_btn_col.addWidget(btn_logo2)
        logo2_btn_col.addWidget(self.lbl_logo2_status)
        logo2_btn_col.addWidget(lbl_logo2_format)
        logo2_btn_col.addStretch()

        logo2_row.addWidget(self.logo2_preview)
        logo2_row.addLayout(logo2_btn_col, 1)

        profile_layout.addLayout(logo2_row)

        # --- Champs texte ---
        def make_field(label_text, placeholder, multiline=False):
            lab = QLabel(label_text)
            lab.setStyleSheet(
                "background:transparent; font-size:12px; font-weight:600;"
            )
            profile_layout.addWidget(lab)

            if multiline:
                field = QTextEdit()
                field.setPlaceholderText(placeholder)
                field.setFixedHeight(70)
            else:
                field = QLineEdit()
                field.setPlaceholderText(placeholder)

            profile_layout.addWidget(field)
            return field

        self.field_nom = make_field(
            "Nom du bureau d'étude", "ex : Silva Environnement"
        )
        if self.bureau_etude:
            self.field_nom.setText(self.bureau_etude)

        self.field_adresse = make_field(
            "Adresse", "ex : Le Champ de la Cure, 58230 Saint Agnan",
            multiline=True
        )
        self.field_tel = make_field("Téléphone", "ex : 06 75 47 29 17")
        self.field_email = make_field(
            "Email", "ex : contact@monbureau.com"
        )
        self.field_web = make_field(
            "Site web", "ex : www.monbureau.com"
        )

        # ==========================================
        # ENREGISTRER / CHARGER LE PROFIL
        # (fait partie du bloc repliable)
        # ==========================================

        profile_io_line = QHBoxLayout()
        profile_io_line.setSpacing(10)

        btn_save_profile = QPushButton("Enregistrer le profil")
        btn_save_profile.setMinimumHeight(44)
        btn_save_profile.clicked.connect(self.save_profile)

        btn_load_profile = QPushButton("Charger un profil")
        btn_load_profile.setMinimumHeight(44)
        btn_load_profile.clicked.connect(self.load_profile)

        profile_io_line.addWidget(btn_save_profile, 1)
        profile_io_line.addWidget(btn_load_profile, 1)

        profile_layout.addLayout(profile_io_line)

        self.lbl_profile_status = QLabel("Aucun profil chargé.")
        self.lbl_profile_status.setWordWrap(True)
        self.lbl_profile_status.setAlignment(Qt.AlignCenter)
        self.lbl_profile_status.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )
        profile_layout.addWidget(self.lbl_profile_status)

        profile_outer.addWidget(self.profile_content)

        main.addWidget(profile_card)

        # ==========================================
        # TITRE DE L'ÉTUDE
        # ==========================================

        titre_card = QFrame()
        titre_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        titre_outer_layout = QVBoxLayout(titre_card)
        titre_outer_layout.setContentsMargins(16, 16, 16, 16)
        titre_outer_layout.setSpacing(10)

        titre_header_row, self.btn_toggle_titre = (
            self._make_collapsible_header("TITRE DE L'ÉTUDE")
        )
        self.btn_toggle_titre.clicked.connect(self.toggle_titre_section)
        titre_outer_layout.addLayout(titre_header_row)

        self.titre_content = QWidget()
        self.titre_content.setStyleSheet("background:transparent;")

        titre_layout = QVBoxLayout(self.titre_content)
        titre_layout.setContentsMargins(0, 0, 0, 0)
        titre_layout.setSpacing(10)

        lab_titre = QLabel("Titre de l'étude")
        lab_titre.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        titre_layout.addWidget(lab_titre)

        self.field_titre_etude = QLineEdit()
        self.field_titre_etude.setPlaceholderText(
            "ex : Étude acoustique des chiroptères en nacelle d'éolienne"
        )
        self.field_titre_etude.setText(
            "Étude acoustique des chiroptères en nacelle d'éolienne"
        )
        titre_layout.addWidget(self.field_titre_etude)

        # --- Parc éolien / Année (pré-remplis depuis le popup de
        # démarrage, modifiables si besoin) ---
        parc_annee_row = QHBoxLayout()
        parc_annee_row.setSpacing(10)

        parc_col = QVBoxLayout()
        lab_parc = QLabel("Parc éolien")
        lab_parc.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        self.field_parc_eolien = QLineEdit()
        self.field_parc_eolien.setPlaceholderText("ex : Vilpion")
        if self.park_name:
            self.field_parc_eolien.setText(self.park_name)
        parc_col.addWidget(lab_parc)
        parc_col.addWidget(self.field_parc_eolien)

        annee_col = QVBoxLayout()
        lab_annee = QLabel("Année")
        lab_annee.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        self.field_annee = QLineEdit()
        self.field_annee.setPlaceholderText("ex : 2025")
        self.field_annee.setText(
            str(self.ref_year) if self.ref_year
            else str(dtmod.date.today().year)
        )
        annee_col.addWidget(lab_annee)
        annee_col.addWidget(self.field_annee)

        parc_annee_row.addLayout(parc_col, 2)
        parc_annee_row.addLayout(annee_col, 1)

        titre_layout.addLayout(parc_annee_row)

        lab_materiel = QLabel(
            "Modèles de matériel — lus automatiquement depuis le "
            f"dossier intégré : {MATERIEL_DIR}\n"
            "(un fichier .txt ou .docx par modèle d'enregistreur, "
            "décrivant le matériel. Optionnel : un fichier "
            "'<Modèle>_figure.docx' à côté, avec juste la photo et "
            "sa légende — inséré juste après la phrase automatique "
            "sur la période de fonctionnement, comme dans le "
            "rapport type)"
        )
        lab_materiel.setWordWrap(True)
        lab_materiel.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        titre_layout.addWidget(lab_materiel)

        materiel_row = QHBoxLayout()
        materiel_row.setSpacing(8)

        btn_materiel_dir = QPushButton("Rafraîchir")
        btn_materiel_dir.clicked.connect(self.pick_materiel_dir)

        self.lbl_materiel_dir = QLabel("Aucun dossier sélectionné")
        self.lbl_materiel_dir.setWordWrap(True)
        self.lbl_materiel_dir.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        materiel_row.addWidget(btn_materiel_dir)
        materiel_row.addWidget(self.lbl_materiel_dir, 1)

        titre_layout.addLayout(materiel_row)

        lab_modele = QLabel("Modèle d'enregistreur utilisé")
        lab_modele.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        titre_layout.addWidget(lab_modele)

        self.combo_modele = QComboBox()
        self.combo_modele.addItem("Aucun (texte générique uniquement)")
        self.combo_modele.setStyleSheet("""
            QComboBox{
                background:black;
                color:white;
                border:1px solid rgba(255,255,255,60);
                border-radius:10px;
                padding:6px 8px;
                font-size:13px;
            }
            QComboBox QAbstractItemView{
                background:#1a1a1a;
                color:white;
                selection-background-color:rgba(90,127,71,220);
                selection-color:white;
                border:1px solid rgba(255,255,255,60);
                outline:0;
            }
        """)
        titre_layout.addWidget(self.combo_modele)

        titre_outer_layout.addWidget(self.titre_content)

        # ==========================================
        # PÉRIODE D'ÉTUDE ET ÉOLIENNE ÉQUIPÉE (repliable)
        # ==========================================

        periode_card = QFrame()
        periode_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        periode_outer_layout = QVBoxLayout(periode_card)
        periode_outer_layout.setContentsMargins(16, 16, 16, 16)
        periode_outer_layout.setSpacing(10)

        periode_header_row, self.btn_toggle_periode = (
            self._make_collapsible_header(
                "PÉRIODE D'ÉTUDE ET ÉOLIENNE ÉQUIPÉE"
            )
        )
        self.btn_toggle_periode.clicked.connect(
            self.toggle_periode_section
        )
        periode_outer_layout.addLayout(periode_header_row)

        self.periode_content = QWidget()
        self.periode_content.setStyleSheet("background:transparent;")

        periode_layout = QVBoxLayout(self.periode_content)
        periode_layout.setContentsMargins(0, 0, 0, 0)
        periode_layout.setSpacing(10)

        lab_periode_intro = QLabel(
            "Alimente la section I.3 du rapport (période d'étude, "
            "éolienne équipée). Le matériel utilisé (Batcorder / "
            "Batlogger / autre) est détecté automatiquement à "
            "partir du modèle d'enregistreur choisi ci-dessus."
        )
        lab_periode_intro.setWordWrap(True)
        lab_periode_intro.setStyleSheet(
            "background:transparent; font-size:11px; color:#cfcfcf;"
        )
        periode_layout.addWidget(lab_periode_intro)

        def make_text_row(label_text, placeholder):
            row = QVBoxLayout()
            row.setSpacing(4)
            lab = QLabel(label_text)
            lab.setStyleSheet(
                "background:transparent; font-size:12px; "
                "font-weight:600;"
            )
            field = QLineEdit()
            field.setPlaceholderText(placeholder)
            row.addWidget(lab)
            row.addWidget(field)
            periode_layout.addLayout(row)
            return field

        self.field_eoliennes = make_text_row(
            "Éolienne(s) équipée(s)", "ex : E3"
        )

        dates_row = QHBoxLayout()
        dates_row.setSpacing(10)

        install_col = QVBoxLayout()
        lab_install = QLabel("Date d'installation")
        lab_install.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        self.field_date_install = QLineEdit()
        self.field_date_install.setPlaceholderText("ex : 26/02/25")
        install_col.addWidget(lab_install)
        install_col.addWidget(self.field_date_install)

        desinstall_col = QVBoxLayout()
        lab_desinstall = QLabel("Date de désinstallation")
        lab_desinstall.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        self.field_date_desinstall = QLineEdit()
        self.field_date_desinstall.setPlaceholderText("ex : 06/11/25")
        desinstall_col.addWidget(lab_desinstall)
        desinstall_col.addWidget(self.field_date_desinstall)

        dates_row.addLayout(install_col)
        dates_row.addLayout(desinstall_col)
        periode_layout.addLayout(dates_row)

        lab_qgis = QLabel(
            "Carte de localisation QGIS — lue automatiquement depuis "
            "le dossier 'Qgis' du projet (déposez-y l'image "
            "exportée depuis QGIS)."
        )
        lab_qgis.setWordWrap(True)
        lab_qgis.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        periode_layout.addWidget(lab_qgis)

        qgis_row = QHBoxLayout()
        qgis_row.setSpacing(8)

        btn_qgis_dir = QPushButton("Rafraîchir")
        btn_qgis_dir.clicked.connect(self.pick_qgis_dir)

        self.lbl_qgis_dir = QLabel(
            f"{self.qgis_dir}\n(dossier du projet, détecté "
            "automatiquement)"
            if self.qgis_dir else "Aucun dossier sélectionné"
        )
        self.lbl_qgis_dir.setWordWrap(True)
        self.lbl_qgis_dir.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        qgis_row.addWidget(btn_qgis_dir)
        qgis_row.addWidget(self.lbl_qgis_dir, 1)

        periode_layout.addLayout(qgis_row)

        periode_outer_layout.addWidget(self.periode_content)

        # ==========================================
        # COMPARAISON AVEC UN SUIVI ANTÉRIEUR (optionnel, repliable)
        # ==========================================

        comparaison_card = QFrame()
        comparaison_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        comparaison_outer_layout = QVBoxLayout(comparaison_card)
        comparaison_outer_layout.setContentsMargins(16, 16, 16, 16)
        comparaison_outer_layout.setSpacing(10)

        comparaison_header_row = QHBoxLayout()
        comparaison_header_row.setSpacing(8)

        self.check_comparaison = QCheckBox()
        self.check_comparaison.setFixedSize(24, 24)
        self.check_comparaison.setStyleSheet("""
            QCheckBox{
                background:transparent;
            }
            QCheckBox::indicator{
                width: 20px;
                height: 20px;
                border: 2px solid rgba(255,255,255,170);
                border-radius: 5px;
                background: rgba(0,0,0,150);
            }
            QCheckBox::indicator:hover{
                border: 2px solid rgba(255,255,255,255);
            }
            QCheckBox::indicator:checked{
                background: rgba(130,184,99,255);
                border: 2px solid rgba(255,255,255,220);
            }
        """)
        self.check_comparaison.toggled.connect(
            self.on_check_comparaison_toggled
        )

        card_title_comparaison = QLabel(
            "COMPARAISON AVEC UN SUIVI ANTÉRIEUR (optionnel)"
        )
        card_title_comparaison.setAlignment(Qt.AlignCenter)
        card_title_comparaison.setStyleSheet("""
            color: #E6E8C8;
            font-weight: 700;
            font-size: 13px;
            letter-spacing: 1px;
            background-color: rgba(0, 0, 0, 0.50);
            border: 1px solid rgba(140,170,90,0.60);
            border-radius: 12px;
            padding: 4px 20px;
        """)

        comparaison_header_row.addWidget(self.check_comparaison)
        comparaison_header_row.addWidget(card_title_comparaison, 1)

        comparaison_outer_layout.addLayout(comparaison_header_row)

        self.comparaison_content = QWidget()
        self.comparaison_content.setStyleSheet("background:transparent;")
        # Replié par défaut : ne se déplie que si la case
        # "Inclure" est cochée (voir on_check_comparaison_toggled).
        self.comparaison_content.setVisible(False)

        comparaison_layout = QVBoxLayout(self.comparaison_content)
        comparaison_layout.setContentsMargins(0, 0, 0, 0)
        comparaison_layout.setSpacing(10)

        lab_comparaison_intro = QLabel(
            "Un suivi antérieur existe sur ce parc : cochez la case "
            "ci-dessus pour ajouter la section III. Comparaison "
            "avec le suivi antérieur au rapport."
        )
        lab_comparaison_intro.setWordWrap(True)
        lab_comparaison_intro.setStyleSheet(
            "background:transparent; font-size:11px; color:#cfcfcf;"
        )
        comparaison_layout.addWidget(lab_comparaison_intro)

        lab_comparaison_annee = QLabel(
            "Année de comparaison (suivi antérieur, si existant)"
        )
        lab_comparaison_annee.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        self.field_annee_suivi_anterieur = QLineEdit()
        self.field_annee_suivi_anterieur.setPlaceholderText(
            "ex : 2021 — laisser vide si aucun suivi antérieur"
        )
        comparaison_layout.addWidget(lab_comparaison_annee)
        comparaison_layout.addWidget(self.field_annee_suivi_anterieur)

        lab_comparaison_graph = QLabel(
            "Graphiques de comparaison — lus automatiquement depuis "
            "le dossier 'Graph' du projet (mots-clés dans le nom du "
            "fichier : Comparaison_total, Comparaison_evolution, "
            "Comparaison_groupes, Comparaison_especes)."
        )
        lab_comparaison_graph.setWordWrap(True)
        lab_comparaison_graph.setStyleSheet(
            "background:transparent; font-size:11px; color:#cfcfcf;"
        )
        comparaison_layout.addWidget(lab_comparaison_graph)

        comparaison_outer_layout.addWidget(self.comparaison_content)

        # ==========================================
        # GRAPHIQUES
        # ==========================================

        graphs_card = QFrame()
        graphs_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        graphs_outer_layout = QVBoxLayout(graphs_card)
        graphs_outer_layout.setContentsMargins(16, 16, 16, 16)
        graphs_outer_layout.setSpacing(10)

        graphs_header_row, self.btn_toggle_graphs = (
            self._make_collapsible_header("GRAPHIQUES")
        )
        self.btn_toggle_graphs.clicked.connect(self.toggle_graphs_section)
        graphs_outer_layout.addLayout(graphs_header_row)

        self.graphs_content = QWidget()
        self.graphs_content.setStyleSheet("background:transparent;")

        graphs_card_layout = QVBoxLayout(self.graphs_content)
        graphs_card_layout.setContentsMargins(0, 0, 0, 0)
        graphs_card_layout.setSpacing(10)

        lab_graphs = QLabel(
            "Dossier des graphiques/tableaux — pointe automatiquement "
            "sur le dossier 'Graph' créé avec le projet."
        )
        lab_graphs.setWordWrap(True)
        lab_graphs.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        graphs_card_layout.addWidget(lab_graphs)

        graphs_row = QHBoxLayout()
        graphs_row.setSpacing(8)

        btn_graphs_dir = QPushButton("Rafraîchir")
        btn_graphs_dir.clicked.connect(self.pick_graphs_dir)

        self.lbl_graphs_dir = QLabel(
            f"{self.graphs_dir}\n(dossier du projet, détecté automatiquement)"
            if self.graphs_dir else "Aucun dossier sélectionné"
        )
        self.lbl_graphs_dir.setWordWrap(True)
        self.lbl_graphs_dir.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        graphs_row.addWidget(btn_graphs_dir)
        graphs_row.addWidget(self.lbl_graphs_dir, 1)

        graphs_card_layout.addLayout(graphs_row)

        lab_scenar = QLabel(
            "Dossier des scénarios de bridage — pointe automatiquement "
            "sur le dossier 'Scenar' créé avec le projet (usage à "
            "venir dans le rapport)."
        )
        lab_scenar.setWordWrap(True)
        lab_scenar.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        graphs_card_layout.addWidget(lab_scenar)

        scenar_row = QHBoxLayout()
        scenar_row.setSpacing(8)

        btn_scenar_dir = QPushButton("Rafraîchir")
        btn_scenar_dir.clicked.connect(self.pick_scenar_dir)

        self.lbl_scenar_dir = QLabel(
            f"{self.scenar_dir}\n(dossier du projet, détecté automatiquement)"
            if self.scenar_dir else "Aucun dossier sélectionné"
        )
        self.lbl_scenar_dir.setWordWrap(True)
        self.lbl_scenar_dir.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        scenar_row.addWidget(btn_scenar_dir)
        scenar_row.addWidget(self.lbl_scenar_dir, 1)

        graphs_card_layout.addLayout(scenar_row)

        graphs_outer_layout.addWidget(self.graphs_content)

        # ==========================================
        # GÉNÉRATION DU RAPPORT TYPE
        # ==========================================

        report_card = QFrame()
        report_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        report_layout = QVBoxLayout(report_card)
        report_layout.setContentsMargins(16, 16, 16, 16)
        report_layout.setSpacing(10)

        report_header_row, self.btn_toggle_report = (
            self._make_collapsible_header("GÉNÉRATION DU RAPPORT")
        )
        self.btn_toggle_report.clicked.connect(self.toggle_report_section)
        report_layout.addLayout(report_header_row)

        self.report_content = QWidget()
        self.report_content.setStyleSheet("background:transparent;")

        report_content_layout = QVBoxLayout(self.report_content)
        report_content_layout.setContentsMargins(0, 0, 0, 0)
        report_content_layout.setSpacing(10)

        note_report = QLabel(
            "Génère une trame Word complète (page de garde avec logo "
            "et coordonnées, sommaire, structure des sections) avec "
            "des repères [À COMPLÉTER] aux endroits où l'analyse doit "
            "être rédigée. Si un dossier de graphiques est indiqué, "
            "les images correspondantes sont insérées automatiquement "
            "à la bonne place ; sinon un repère à insérer manuellement "
            "est laissé."
        )
        note_report.setWordWrap(True)
        note_report.setStyleSheet(
            "color:#f2e6d0; font-size:10px; background:transparent;"
        )
        report_content_layout.addWidget(note_report)

        btn_generate = QPushButton("GÉNÉRER UN RAPPORT TYPE")
        btn_generate.setMinimumHeight(50)
        btn_generate.clicked.connect(self.generate_report_template)
        report_content_layout.addWidget(btn_generate)

        report_layout.addWidget(self.report_content)

        # ==========================================
        # STYLE PERSONNALISÉ DU RAPPORT (repliable)
        # ==========================================

        style_card = QFrame()
        style_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        style_outer = QVBoxLayout(style_card)
        style_outer.setContentsMargins(16, 16, 16, 16)
        style_outer.setSpacing(10)

        style_header_row, self.btn_toggle_style = (
            self._make_collapsible_header("STYLE PERSONNALISÉ DU RAPPORT")
        )
        self.btn_toggle_style.clicked.connect(self.toggle_style_section)
        style_outer.addLayout(style_header_row)

        self.style_content = QWidget()
        self.style_content.setStyleSheet("background:transparent;")

        style_layout = QVBoxLayout(self.style_content)
        style_layout.setContentsMargins(0, 0, 0, 0)
        style_layout.setSpacing(10)

        lab_font = QLabel("Police d'écriture du rapport")
        lab_font.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        style_layout.addWidget(lab_font)

        self.font_combo = QFontComboBox()
        self.font_combo.setCurrentFont(QFont("Century Gothic"))
        self.font_combo.setStyleSheet("""
            QFontComboBox{
                background:black;
                color:white;
                border:1px solid rgba(255,255,255,60);
                border-radius:10px;
                padding:6px 8px;
                font-size:13px;
            }
        """)
        style_layout.addWidget(self.font_combo)

        # --- Tailles de police ---
        lab_sizes = QLabel("Tailles de police")
        lab_sizes.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        style_layout.addWidget(lab_sizes)

        def make_size_row(label_text, default_value):
            row = QHBoxLayout()
            row.setSpacing(8)

            lab = QLabel(label_text)
            lab.setStyleSheet(
                "background:transparent; font-size:11px;"
            )
            lab.setFixedWidth(180)

            spin = QSpinBox()
            spin.setRange(6, 48)
            spin.setValue(default_value)
            spin.setSuffix(" pt")
            spin.setStyleSheet("""
                QSpinBox{
                    background:black;
                    color:white;
                    border:1px solid rgba(255,255,255,60);
                    border-radius:8px;
                    padding:4px 6px;
                    font-size:12px;
                }
            """)

            row.addWidget(lab)
            row.addWidget(spin)
            row.addStretch()

            style_layout.addLayout(row)
            return spin

        self.spin_size_titre = make_size_row(
            "Titres (Sommaire, I., II....)", 18
        )
        self.spin_size_soustitre = make_size_row(
            "Sous-titres (I.1, Contexte...)", 16
        )
        self.spin_size_corps = make_size_row(
            "Texte d'analyse (corps)", 11
        )

        # --- Couleurs ---
        lab_colors = QLabel("Couleurs")
        lab_colors.setStyleSheet(
            "background:transparent; font-size:12px; font-weight:600;"
        )
        style_layout.addWidget(lab_colors)

        self.color_titre = QColor(0x2E, 0x3A, 0x4A)
        self.color_corps = QColor(0x00, 0x00, 0x00)

        def make_color_row(label_text, initial_color):
            row = QHBoxLayout()
            row.setSpacing(8)

            lab = QLabel(label_text)
            lab.setStyleSheet(
                "background:transparent; font-size:11px;"
            )
            lab.setFixedWidth(180)

            swatch = QPushButton()
            swatch.setFixedSize(32, 24)
            swatch.setStyleSheet(
                f"background:{initial_color.name()}; "
                f"border:1px solid rgba(255,255,255,80); "
                f"border-radius:6px;"
            )

            row.addWidget(lab)
            row.addWidget(swatch)
            row.addStretch()

            style_layout.addLayout(row)
            return swatch

        self.swatch_titre = make_color_row(
            "Couleur des titres", self.color_titre
        )
        self.swatch_titre.clicked.connect(
            lambda: self._pick_color("titre")
        )

        self.swatch_corps = make_color_row(
            "Couleur du texte", self.color_corps
        )
        self.swatch_corps.clicked.connect(
            lambda: self._pick_color("corps")
        )

        # --- Enregistrer / charger le style ---
        style_io_line = QHBoxLayout()
        style_io_line.setSpacing(10)

        btn_save_style = QPushButton("Enregistrer le style")
        btn_save_style.setMinimumHeight(40)
        btn_save_style.clicked.connect(self.save_style)

        btn_load_style = QPushButton("Charger un style")
        btn_load_style.setMinimumHeight(40)
        btn_load_style.clicked.connect(self.load_style)

        style_io_line.addWidget(btn_save_style, 1)
        style_io_line.addWidget(btn_load_style, 1)

        style_layout.addLayout(style_io_line)

        style_outer.addWidget(self.style_content)

        main.addWidget(style_card)

        main.addWidget(titre_card)

        main.addWidget(periode_card)

        main.addWidget(comparaison_card)

        main.addWidget(graphs_card)

        # ==========================================
        # BIBLIOGRAPHIE (repliable)
        # ==========================================

        biblio_card = QFrame()
        biblio_card.setStyleSheet("""
            QFrame{
                background:rgba(59,41,25,190);
                border:1px solid rgba(255,255,255,40);
                border-radius:16px;
            }
        """)

        biblio_outer = QVBoxLayout(biblio_card)
        biblio_outer.setContentsMargins(16, 16, 16, 16)
        biblio_outer.setSpacing(10)

        biblio_header_row, self.btn_toggle_biblio = (
            self._make_collapsible_header("BIBLIOGRAPHIE")
        )
        self.btn_toggle_biblio.clicked.connect(self.toggle_biblio_section)
        biblio_outer.addLayout(biblio_header_row)

        self.biblio_content = QWidget()
        self.biblio_content.setStyleSheet("background:transparent;")

        biblio_layout = QVBoxLayout(self.biblio_content)
        biblio_layout.setContentsMargins(0, 0, 0, 0)
        biblio_layout.setSpacing(10)

        lab_biblio = QLabel(
            "Bibliographie — lue automatiquement depuis le dossier "
            f"intégré : {BIBLIO_DIR}\n"
            "(un ou plusieurs fichiers .docx/.txt, une référence par "
            "ligne/paragraphe — vous pouvez l'enrichir au fil du "
            "temps, seules les références cochées ci-dessous seront "
            "ajoutées au rapport)."
        )
        lab_biblio.setWordWrap(True)
        lab_biblio.setStyleSheet(
            "background:transparent; font-size:11px; color:#f2e6d0;"
        )
        biblio_layout.addWidget(lab_biblio)

        biblio_btn_row = QHBoxLayout()
        biblio_btn_row.setSpacing(8)

        btn_biblio_file = QPushButton("Rafraîchir")
        btn_biblio_file.clicked.connect(self.pick_biblio_file)

        self.lbl_biblio_file = QLabel("Aucun fichier sélectionné")
        self.lbl_biblio_file.setWordWrap(True)
        self.lbl_biblio_file.setStyleSheet(
            "background:transparent; color:#999999; font-size:11px;"
        )

        biblio_btn_row.addWidget(btn_biblio_file)
        biblio_btn_row.addWidget(self.lbl_biblio_file, 1)

        biblio_layout.addLayout(biblio_btn_row)

        biblio_select_row = QHBoxLayout()
        biblio_select_row.setSpacing(8)

        btn_biblio_all = QPushButton("Tout cocher")
        btn_biblio_all.clicked.connect(
            lambda: self._set_all_biblio_checks(True)
        )

        btn_biblio_none = QPushButton("Tout décocher")
        btn_biblio_none.clicked.connect(
            lambda: self._set_all_biblio_checks(False)
        )

        biblio_select_row.addWidget(btn_biblio_all)
        biblio_select_row.addWidget(btn_biblio_none)

        biblio_layout.addLayout(biblio_select_row)

        self.biblio_list = QListWidget()
        self.biblio_list.setStyleSheet("""
            QListWidget{
                background:black;
                color:white;
                border:1px solid rgba(255,255,255,60);
                border-radius:10px;
                padding:6px;
                font-size:12px;
            }
            QListWidget::item{
                padding:5px 2px;
            }
            QListWidget::indicator{
                width:16px;
                height:16px;
                border:1px solid rgba(255,255,255,120);
                border-radius:4px;
                background:rgba(255,255,255,20);
            }
            QListWidget::indicator:hover{
                border:1px solid rgba(255,255,255,200);
            }
            QListWidget::indicator:checked{
                background:rgba(160,210,120,255);
                border:1px solid rgba(255,255,255,180);
            }
        """)
        self.biblio_list.setMaximumHeight(220)
        biblio_layout.addWidget(self.biblio_list)

        biblio_outer.addWidget(self.biblio_content)

        main.addWidget(biblio_card)

        main.addWidget(report_card)

        main.addStretch()

        scroll.setWidget(content)
        outer.addWidget(scroll)

    # ======================================================

    def toggle_profile_section(self):

        visible = self.profile_content.isVisible()

        self.profile_content.setVisible(not visible)

        self.btn_toggle_profile.setText("▶" if visible else "▼")

    # ======================================================

    def toggle_style_section(self):

        visible = self.style_content.isVisible()

        self.style_content.setVisible(not visible)

        self.btn_toggle_style.setText("▶" if visible else "▼")

    # ======================================================

    def toggle_biblio_section(self):

        visible = self.biblio_content.isVisible()

        self.biblio_content.setVisible(not visible)

        self.btn_toggle_biblio.setText("▶" if visible else "▼")

    # ======================================================

    def refresh_all(self):
        """
        Relit tous les dossiers/fichiers intégrés au projet en une
        fois (matériel, QGIS, graphiques, Scenar, bibliographie) —
        équivalent à cliquer sur chacun des boutons "Rafraîchir"
        séparément. Appelée automatiquement à chaque fois que
        l'onglet Rapport devient l'onglet actif (voir Suite dans le
        launcher), pour ne pas avoir à le faire manuellement.
        """

        self.pick_materiel_dir()
        self.pick_qgis_dir()
        self.pick_graphs_dir()
        self.pick_scenar_dir()
        self.pick_biblio_file()

    # ======================================================

    def toggle_titre_section(self):

        visible = self.titre_content.isVisible()

        self.titre_content.setVisible(not visible)

        self.btn_toggle_titre.setText("▶" if visible else "▼")

    # ======================================================

    def toggle_periode_section(self):

        visible = self.periode_content.isVisible()

        self.periode_content.setVisible(not visible)

        self.btn_toggle_periode.setText("▶" if visible else "▼")

    # ======================================================

    def on_check_comparaison_toggled(self, checked):
        """
        La case à cocher pilote seule le repli/dépli du bloc (plus
        de flèche séparée) : cochée -> déplié, décochée -> replié.
        """

        self.comparaison_content.setVisible(checked)

    # ======================================================

    def toggle_graphs_section(self):

        visible = self.graphs_content.isVisible()

        self.graphs_content.setVisible(not visible)

        self.btn_toggle_graphs.setText("▶" if visible else "▼")

    # ======================================================

    def toggle_report_section(self):

        visible = self.report_content.isVisible()

        self.report_content.setVisible(not visible)

        self.btn_toggle_report.setText("▶" if visible else "▼")

    # ======================================================

    def pick_biblio_file(self):
        """
        Relit le dossier intégré de bibliographie
        (dossiers pour rapport/bibliographie) et met à jour la
        liste des références. Les cases déjà cochées avant un
        rafraîchissement sont conservées si la référence existe
        toujours.
        """

        # Mémorise les références déjà cochées, pour les conserver
        # après le rafraîchissement.
        previously_checked = {
            self.biblio_list.item(i).text()
            for i in range(self.biblio_list.count())
            if self.biblio_list.item(i).checkState() == Qt.Checked
        }

        if not os.path.isdir(BIBLIO_DIR):
            self.biblio_list.clear()
            self.lbl_biblio_file.setText(
                f"Dossier introuvable : {BIBLIO_DIR}\n"
                f"(créez-le à côté du dossier 'modules' pour "
                f"y déposer vos fichiers de bibliographie)"
            )
            return

        entries = []

        try:
            for f in sorted(os.listdir(BIBLIO_DIR)):
                if f.startswith("~$") or f.startswith("."):
                    continue
                if f.lower().endswith((".txt", ".docx")):
                    entries.extend(
                        self._read_biblio_entries(
                            os.path.join(BIBLIO_DIR, f)
                        )
                    )
        except Exception as e:
            QMessageBox.critical(
                self, "Erreur",
                f"Impossible de lire le dossier de bibliographie :\n{e}"
            )
            return

        self.biblio_list.clear()

        for entry in entries:
            item = QListWidgetItem(entry)
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(
                Qt.Checked if entry in previously_checked else Qt.Unchecked
            )
            self.biblio_list.addItem(item)

        self.lbl_biblio_file.setText(
            f"{BIBLIO_DIR}\n({len(entries)} référence(s) trouvée(s))"
        )

    # ======================================================

    def _read_biblio_entries(self, path):
        """Lit un fichier de bibliographie (.docx ou .txt) et
        retourne la liste des références (une par ligne/paragraphe
        non vide)."""

        entries = []

        if path.lower().endswith(".docx"):
            src = Document(path)
            for p in src.paragraphs:
                if p.text.strip():
                    entries.append(p.text.strip())
        else:
            with open(path, "r", encoding="utf-8") as f:
                for ligne in f.read().split("\n"):
                    if ligne.strip():
                        entries.append(ligne.strip())

        return entries

    # ======================================================

    def _set_all_biblio_checks(self, checked):

        state = Qt.Checked if checked else Qt.Unchecked

        for i in range(self.biblio_list.count()):
            self.biblio_list.item(i).setCheckState(state)

    # ======================================================

    def _pick_color(self, which):

        current = (
            self.color_titre if which == "titre" else self.color_corps
        )

        color = QColorDialog.getColor(
            current, self, "Choisir une couleur"
        )

        if not color.isValid():
            return

        if which == "titre":
            self.color_titre = color
            swatch = self.swatch_titre
        else:
            self.color_corps = color
            swatch = self.swatch_corps

        swatch.setStyleSheet(
            f"background:{color.name()}; "
            f"border:1px solid rgba(255,255,255,80); "
            f"border-radius:6px;"
        )

    # ======================================================

    def save_style_to_path(self, out_path):
        """Enregistre le style directement vers un chemin donné,
        sans boîte de dialogue (utilisé par la sauvegarde globale
        du projet)."""

        data = {
            "police": self.font_combo.currentFont().family(),
            "taille_titre": self.spin_size_titre.value(),
            "taille_soustitre": self.spin_size_soustitre.value(),
            "taille_corps": self.spin_size_corps.value(),
            "couleur_titre": self.color_titre.name(),
            "couleur_corps": self.color_corps.name(),
        }

        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    # ======================================================

    def save_style(self):
        """Enregistre le style personnalisé (police, tailles,
        couleurs) dans un fichier .json réutilisable."""

        try:

            out_path, _ = QFileDialog.getSaveFileName(
                self,
                "Enregistrer le style du rapport",
                "style_rapport.json",
                "Style (*.json)"
            )

            if not out_path:
                return

            if not out_path.lower().endswith(".json"):
                out_path += ".json"

            self.save_style_to_path(out_path)

            QMessageBox.information(
                self, "Succès",
                f"Style enregistré :\n{out_path}"
            )

        except Exception as e:
            QMessageBox.critical(
                self, "Erreur",
                f"Impossible d'enregistrer le style :\n{e}"
            )

    # ======================================================

    def load_style_from_path(self, in_path):
        """Charge un style directement depuis un chemin donné, sans
        boîte de dialogue (utilisé par le chargement global du
        projet)."""

        with open(in_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        police = data.get("police")
        if police:
            self.font_combo.setCurrentFont(QFont(police))
        else:
            self.font_combo.setCurrentFont(QFont("Century Gothic"))

        self.spin_size_titre.setValue(data.get("taille_titre", 18))
        self.spin_size_soustitre.setValue(
            data.get("taille_soustitre", 16)
        )
        self.spin_size_corps.setValue(data.get("taille_corps", 11))

        couleur_titre = data.get("couleur_titre")
        if couleur_titre:
            self.color_titre = QColor(couleur_titre)
            self.swatch_titre.setStyleSheet(
                f"background:{self.color_titre.name()}; "
                f"border:1px solid rgba(255,255,255,80); "
                f"border-radius:6px;"
            )

        couleur_corps = data.get("couleur_corps")
        if couleur_corps:
            self.color_corps = QColor(couleur_corps)
            self.swatch_corps.setStyleSheet(
                f"background:{self.color_corps.name()}; "
                f"border:1px solid rgba(255,255,255,80); "
                f"border-radius:6px;"
            )

        # Un style vient d'être chargé (projet existant ou style
        # rechargé manuellement) : on replie le bloc, il n'y a pas
        # besoin de le garder ouvert (toujours possible de le
        # rouvrir pour vérifier/modifier).
        self.style_content.setVisible(False)
        self.btn_toggle_style.setText("▶")

    # ======================================================

    def load_style(self):
        """Charge un style personnalisé précédemment enregistré."""

        try:

            in_path, _ = QFileDialog.getOpenFileName(
                self,
                "Charger un style de rapport",
                "",
                "Style (*.json)"
            )

            if not in_path:
                return

            self.load_style_from_path(in_path)

            QMessageBox.information(
                self, "Succès",
                f"Style chargé depuis :\n{in_path}"
            )

        except Exception as e:
            QMessageBox.critical(
                self, "Erreur",
                f"Impossible de charger le style :\n{e}"
            )

    # ======================================================

    def pick_graphs_dir(self):
        """Actualise l'affichage du dossier 'Graph' du projet
        (lié automatiquement à la création du projet)."""

        if not self.graphs_dir or not os.path.isdir(self.graphs_dir):
            self.lbl_graphs_dir.setText(
                "Aucun dossier lié — le dossier 'Graph' du projet "
                "n'a pas été trouvé."
            )
            return

        images = [
            f for f in os.listdir(self.graphs_dir)
            if f.lower().endswith((".jpg", ".jpeg", ".png"))
        ]

        self.lbl_graphs_dir.setText(
            f"{self.graphs_dir}\n({len(images)} image(s) trouvée(s))"
        )

    # ======================================================

    def pick_scenar_dir(self):
        """Actualise l'affichage du dossier 'Scenar' du projet
        (lié automatiquement à la création du projet)."""

        if not self.scenar_dir or not os.path.isdir(self.scenar_dir):
            self.lbl_scenar_dir.setText(
                "Aucun dossier lié — le dossier 'Scenar' du projet "
                "n'a pas été trouvé."
            )
            return

        fichiers = [
            f for f in os.listdir(self.scenar_dir)
            if not f.startswith(".")
        ]

        self.lbl_scenar_dir.setText(
            f"{self.scenar_dir}\n({len(fichiers)} fichier(s) trouvé(s))"
        )

    # ======================================================

    def pick_qgis_dir(self):
        """Actualise l'affichage du dossier 'Qgis' du projet
        (lié automatiquement à la création du projet), qui doit
        contenir la carte de localisation exportée depuis QGIS."""

        if not self.qgis_dir or not os.path.isdir(self.qgis_dir):
            self.lbl_qgis_dir.setText(
                "Aucun dossier lié — le dossier 'Qgis' du projet "
                "n'a pas été trouvé."
            )
            return

        images = [
            f for f in os.listdir(self.qgis_dir)
            if f.lower().endswith((".jpg", ".jpeg", ".png"))
            and not f.startswith(".")
        ]

        self.lbl_qgis_dir.setText(
            f"{self.qgis_dir}\n({len(images)} image(s) trouvée(s))"
        )

    # ======================================================

    def pick_materiel_dir(self):
        """
        Relit le dossier intégré des modèles de matériel
        (dossiers pour rapport/materiel utilisé) et met à jour le
        menu déroulant en conséquence.

        Convention pour séparer texte et image (utile pour que la
        phrase automatique 'L'enregistreur a fonctionné...' se
        retrouve entre le texte descriptif et la photo, comme dans
        le rapport type) : un fichier compagnon
        '<Modèle>_figure.jpg/.png' à côté de '<Modèle>.docx' est
        traité comme l'image à insérer après la phrase automatique.
        Sa légende est cherchée en priorité dans un fichier
        '<Modèle>_figure.docx' de même nom (texte seul, voir
        _resolve_caption) ; à défaut, un texte générique est utilisé.
        Nom simplifié aussi accepté, sans préfixe de modèle :
        'figure.jpg' + 'figure.docx' (légende), utilisé par défaut
        quel que soit le modèle sélectionné — pratique quand un
        seul type de matériel est utilisé.
        Ancien format encore supporté en repli : si aucune image
        '<Modèle>_figure.jpg/.png' n'existe, un '<Modèle>_figure.docx'
        contenant à la fois l'image et son texte est inséré tel
        quel (sans légende numérotée automatique dans ce cas).
        Ce fichier compagnon n'apparaît pas comme un modèle séparé
        dans le menu déroulant.
        """

        path = MATERIEL_DIR

        self.materiel_dir = path
        self.materiel_files = {}
        self.materiel_figure_files = {}

        if not os.path.isdir(path):
            self.lbl_materiel_dir.setText(
                f"Dossier introuvable : {path}\n"
                f"(créez-le à côté du dossier 'modules' pour "
                f"y déposer vos descriptions de matériel)"
            )
        else:
            for f in sorted(os.listdir(path)):
                if f.startswith("~$") or f.startswith("."):
                    continue

                ext = os.path.splitext(f)[1].lower()

                if ext not in (".txt", ".docx", ".jpg", ".jpeg", ".png"):
                    continue

                base = os.path.splitext(f)[0]
                base_low = base.lower()

                is_figure_file = (
                    base_low.endswith("_figure") or base_low == "figure"
                )

                if is_figure_file:
                    nom_modele = (
                        "__default__" if base_low == "figure"
                        else base[: -len("_figure")]
                    )
                    if ext in (".jpg", ".jpeg", ".png"):
                        # Nouvelle convention (image + légende
                        # séparée) : priorité sur l'ancien format.
                        self.materiel_figure_files[nom_modele] = (
                            os.path.join(path, f)
                        )
                    elif ext == ".docx" and (
                        nom_modele not in self.materiel_figure_files
                    ):
                        # Ancien format (image + texte dans le même
                        # .docx), gardé en repli seulement si aucune
                        # image séparée n'existe déjà pour ce modèle.
                        self.materiel_figure_files[nom_modele] = (
                            os.path.join(path, f)
                        )
                elif ext in (".txt", ".docx"):
                    nom_modele = base
                    self.materiel_files[nom_modele] = os.path.join(
                        path, f
                    )

            self.lbl_materiel_dir.setText(
                f"{path}\n"
                f"({len(self.materiel_files)} modèle(s) trouvé(s))"
            )

        self.combo_modele.clear()
        self.combo_modele.addItem("Aucun (texte générique uniquement)")

        for nom_modele in self.materiel_files:
            self.combo_modele.addItem(nom_modele)

    # ======================================================

    def _find_statuts_file(self):
        """
        Cherche, dans le dossier intégré des statuts d'espèces
        (dossiers pour rapport/Statuts des espèces), le fichier
        unique (.docx ou .txt) à insérer tel quel en II.3. Pas de
        choix pour l'utilisateur : un seul fichier fixe attendu
        (ex : "Statuts.docx"). Ignore les fichiers verrous (~$...)
        et cachés.
        """

        path = STATUTS_DIR

        if not os.path.isdir(path):
            return None

        for f in sorted(os.listdir(path)):
            if f.startswith("~$") or f.startswith("."):
                continue
            if f.lower().endswith((".txt", ".docx")):
                return os.path.join(path, f)

        return None

    # ======================================================

    def _find_statuts_part_file(self, keyword):
        """
        Cherche, dans le dossier intégré des statuts d'espèces
        (dossiers pour rapport/Statuts des espèces), le fichier
        (.docx ou .txt) dont le nom contient le mot-clé donné (ex :
        "part1", "part2"...), insensible à la casse. Permet de
        découper la II.3 en plusieurs blocs de texte, entrecoupés
        de tableaux/figures (voir génération du rapport). Ignore
        les fichiers verrous (~$...) et cachés.
        """

        path = STATUTS_DIR

        if not os.path.isdir(path):
            return None

        for f in sorted(os.listdir(path)):
            if f.startswith("~$") or f.startswith("."):
                continue
            if not f.lower().endswith((".txt", ".docx")):
                continue
            if keyword.lower() in f.lower():
                return os.path.join(path, f)

        return None

    # ======================================================

    def _find_analyse_file(self):
        """
        Cherche, dans le dossier intégré de la méthode d'analyse
        des enregistrements (dossiers pour rapport/Analyse des
        enregistrements), le fichier unique (.docx ou .txt) à
        insérer tel quel en I.2. Pas de choix pour l'utilisateur :
        un seul fichier fixe attendu, propre au bureau d'étude
        (logiciels utilisés, méthode de quantification de
        l'activité...). Ignore les fichiers verrous (~$...) et
        cachés.
        """

        path = ANALYSE_DIR

        if not os.path.isdir(path):
            return None

        for f in sorted(os.listdir(path)):
            if f.startswith("~$") or f.startswith("."):
                continue
            if f.lower().endswith((".txt", ".docx")):
                return os.path.join(path, f)

        return None

    # ------------------------------------------------------

    def _load_resume_rapport(self):
        """
        Lit 'resume_rapport.json', écrit par Graph (bouton export
        des tableaux) dans le dossier 'Graph' du projet, qui
        contient les chiffres bruts calculés depuis le fichier
        source (nombre total de contacts, nombre de nuits
        positives, mois couverts, dates de premier/dernier
        contact). Permet d'auto-remplir certaines phrases du
        rapport (ex : II.1) sans recalculer quoi que ce soit côté
        Rapport. Retourne None si absent/illisible — auto-remplissage
        alors non disponible, l'utilisateur reste libre de rédiger
        à la main.
        """

        if not self.graphs_dir or not os.path.isdir(self.graphs_dir):
            return None

        path = os.path.join(self.graphs_dir, "resume_rapport.json")

        if not os.path.exists(path):
            return None

        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None

    # ------------------------------------------------------

    def _parse_date_fr(self, text):
        """
        Tente de convertir un texte de date saisi à la main
        (formats courants JJ/MM/AA, JJ/MM/AAAA, JJ-MM-AAAA) en
        objet date. Retourne None si le texte est vide ou
        illisible — pas d'erreur bloquante, l'auto-remplissage
        retombe simplement sur une estimation alternative.
        """

        if not text:
            return None

        for fmt in ("%d/%m/%Y", "%d/%m/%y", "%d-%m-%Y", "%d-%m-%y"):
            try:
                return dtmod.datetime.strptime(text, fmt).date()
            except ValueError:
                continue

        return None

    # ------------------------------------------------------

    def pick_logo(self):

        path, _ = QFileDialog.getOpenFileName(
            self,
            "Choisir le logo du bureau d'étude",
            "",
            "Images (*.png *.jpg *.jpeg)"
        )

        if not path:
            return

        self.logo_path = path

        pix = QPixmap(path)

        if pix.isNull():
            self.lbl_logo_status.setText("Fichier image invalide.")
            self.logo_path = None
            return

        self.logo_preview.setPixmap(
            pix.scaled(
                self.logo_preview.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
        )

        self.lbl_logo_status.setText(os.path.basename(path))

    # ======================================================

    def pick_logo_secondaire(self):

        path, _ = QFileDialog.getOpenFileName(
            self,
            "Choisir le logo secondaire",
            "",
            "Images (*.png *.jpg *.jpeg)"
        )

        if not path:
            return

        self.logo_secondaire_path = path

        pix = QPixmap(path)

        if pix.isNull():
            self.lbl_logo2_status.setText("Fichier image invalide.")
            self.logo_secondaire_path = None
            return

        self.logo2_preview.setPixmap(
            pix.scaled(
                self.logo2_preview.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
        )

        self.lbl_logo2_status.setText(os.path.basename(path))

    # ======================================================

    def save_profile_to_path(self, profile_path):
        """Enregistre le profil directement vers un chemin donné,
        sans boîte de dialogue (utilisé par la sauvegarde globale
        du projet). Retourne (logo_filename, logo2_filename)."""

        nom = self.field_nom.text().strip()

        if not nom:
            raise Exception("Le nom du bureau d'étude est obligatoire.")

        out_dir = os.path.dirname(profile_path)
        base_name = os.path.splitext(
            os.path.basename(profile_path)
        )[0]

        def _copy_logo(src_path, suffix):
            if not src_path or not os.path.exists(src_path):
                return None
            ext = os.path.splitext(src_path)[1] or ".png"
            filename = f"{base_name}_{suffix}{ext}"
            dest_path = os.path.join(out_dir, filename)
            same_file = (
                os.path.exists(dest_path)
                and os.path.samefile(src_path, dest_path)
            )
            if not same_file:
                shutil.copy(src_path, dest_path)
            return filename, dest_path

        logo_filename = None
        result = _copy_logo(self.logo_path, "logo")
        if result:
            logo_filename, self.logo_path = result

        logo2_filename = None
        result2 = _copy_logo(self.logo_secondaire_path, "logo2")
        if result2:
            logo2_filename, self.logo_secondaire_path = result2

        data = {
            "nom": nom,
            "adresse": self.field_adresse.toPlainText().strip(),
            "telephone": self.field_tel.text().strip(),
            "email": self.field_email.text().strip(),
            "site_web": self.field_web.text().strip(),
            "logo_fichier": logo_filename,
            "logo_secondaire_fichier": logo2_filename,
        }

        with open(profile_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        self.profile_path = profile_path

        self.lbl_profile_status.setText(
            f"Profil enregistré : {profile_path}"
        )

        return logo_filename, logo2_filename

    # ======================================================

    def save_profile(self):
        """
        Enregistre le profil (coordonnées + logo) : un fichier .json
        nommé au choix, accompagné d'une copie du logo dans le même
        dossier, pour rester portable et rechargeable.
        """

        try:

            nom = self.field_nom.text().strip()

            if not nom:
                raise Exception(
                    "Le nom du bureau d'étude est obligatoire."
                )

            safe_nom = "".join(
                c if c.isalnum() or c in " _-" else "_"
                for c in nom
            ).strip() or "profil_bureau"

            profile_path, _ = QFileDialog.getSaveFileName(
                self,
                "Enregistrer le profil du bureau d'étude",
                f"{safe_nom}.json",
                "Profil (*.json)"
            )

            if not profile_path:
                return

            if not profile_path.lower().endswith(".json"):
                profile_path += ".json"

            logo_filename, logo2_filename = self.save_profile_to_path(
                profile_path
            )

            QMessageBox.information(
                self, "Succès",
                f"Profil du bureau enregistré :\n{profile_path}"
                + (f"\n+ {logo_filename}" if logo_filename else "")
                + (f"\n+ {logo2_filename}" if logo2_filename else "")
            )

        except Exception as e:
            QMessageBox.critical(
                self, "Erreur",
                f"Impossible d'enregistrer le profil :\n{e}"
            )

    # ======================================================

    def load_profile_from_path(self, in_path):
        """Charge un profil directement depuis un chemin donné,
        sans boîte de dialogue (utilisé par le chargement global
        du projet)."""

        with open(in_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        self.field_nom.setText(data.get("nom", ""))
        self.field_adresse.setPlainText(data.get("adresse", ""))
        self.field_tel.setText(data.get("telephone", ""))
        self.field_email.setText(data.get("email", ""))
        self.field_web.setText(data.get("site_web", ""))

        def _load_logo(filename, preview_widget, status_widget):
            if not filename:
                return None
            full_path = os.path.join(
                os.path.dirname(in_path), filename
            )
            if os.path.exists(full_path):
                pix = QPixmap(full_path)
                if not pix.isNull():
                    preview_widget.setPixmap(
                        pix.scaled(
                            preview_widget.size(),
                            Qt.KeepAspectRatio,
                            Qt.SmoothTransformation
                        )
                    )
                    status_widget.setText(os.path.basename(full_path))
                    return full_path
            else:
                preview_widget.setText("Logo introuvable")
                status_widget.setText(
                    f"Fichier logo attendu introuvable : {filename}"
                )
            return None

        self.logo_path = _load_logo(
            data.get("logo_fichier"),
            self.logo_preview,
            self.lbl_logo_status
        )

        self.logo_secondaire_path = _load_logo(
            data.get("logo_secondaire_fichier"),
            self.logo2_preview,
            self.lbl_logo2_status
        )

        self.profile_path = in_path

        self.lbl_profile_status.setText(
            f"Profil chargé : {in_path}"
        )

        # Bloc déjà rempli par le chargement : on le replie pour
        # ne pas encombrer l'écran (toujours possible de le
        # rouvrir pour vérifier/modifier).
        if self.field_nom.text().strip() or self.field_adresse.toPlainText().strip():
            self.profile_content.setVisible(False)
            self.btn_toggle_profile.setText("▶")

    # ======================================================

    def load_profile(self):
        """
        Charge un profil précédemment enregistré (profil_bureau.json
        + logo associé dans le même dossier).
        """

        try:

            in_path, _ = QFileDialog.getOpenFileName(
                self,
                "Charger un profil de bureau d'étude",
                "",
                "Profil (*.json)"
            )

            if not in_path:
                return

            self.load_profile_from_path(in_path)

            QMessageBox.information(
                self, "Succès",
                f"Profil chargé depuis :\n{in_path}"
            )

        except Exception as e:
            QMessageBox.critical(
                self, "Erreur",
                f"Impossible de charger le profil :\n{e}"
            )

    # ======================================================
    # GÉNÉRATION DU RAPPORT TYPE
    # ======================================================

    def _add_toc_field(self, doc):
        """Insère un champ Sommaire (TOC) Word natif — à mettre à
        jour dans Word via clic droit → 'Mettre à jour les champs'."""

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = (
            "Clic droit ici puis « Mettre à jour les champs » "
            "pour générer le sommaire."
        )

        fld_sep.append(placeholder)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(fld_end)

    # ------------------------------------------------------

    def _add_list_field(self, doc, label):
        """
        Insère un champ Word natif "Liste des figures"/"Liste des
        tableaux" (TOC \\c "<label>"), qui compile automatiquement
        toutes les légendes numérotées via _add_numbered_caption
        portant ce même label, avec leur numéro de page — même
        principe que le sommaire (_add_toc_field), à mettre à jour
        dans Word via clic droit → 'Mettre à jour les champs'.
        """

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f'TOC \\h \\z \\c "{label}"'

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = (
            "Clic droit ici puis « Mettre à jour les champs » "
            "pour générer la liste."
        )
        fld_sep.append(placeholder)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(fld_end)

    # ------------------------------------------------------

    def _add_numbered_caption(self, doc, label, text):
        """
        Insère une légende numérotée automatiquement par Word (
        équivalent d'Insertion > Légende) : "<label> <N> : texte",
        où <N> est un champ SEQ recalculé par Word — pas un simple
        numéro incrémenté ici en Python, ce qui serait faux dès
        qu'une figure est ajoutée/retirée manuellement après coup
        dans Word. Ce même champ SEQ est ce qui permet à
        _add_list_field de générer automatiquement la liste des
        figures/tableaux avec numéros de page. Le numéro affiché
        tant que les champs n'ont pas été mis à jour dans Word est
        provisoire ("1") — comme pour le sommaire.
        """

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        # Empêche Word d'insérer un saut de page entre la légende
        # et le paragraphe qui suit (l'image/le tableau) : les deux
        # doivent toujours rester ensemble sur la même page.
        p.paragraph_format.keep_with_next = True

        def _add_caption_run(text_run=None, is_field=False):
            run = p.add_run(text_run or "")
            run.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            return run

        _add_caption_run(f"{label} ")

        run_seq = _add_caption_run()
        r = run_seq._r

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" SEQ {label} \\* ARABIC "

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = "1"
        fld_sep.append(placeholder)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(fld_end)

        _add_caption_run(f" : {text}")

        return p

    # ------------------------------------------------------

    def _add_todo(self, doc, text="Analyse à rédiger."):
        """Ajoute un paragraphe repère [À COMPLÉTER] bien visible."""

        p = doc.add_paragraph()
        run = p.add_run(f"[À COMPLÉTER — {text}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
        run.font.size = Pt(self.spin_size_corps.value())
        return p

    # ------------------------------------------------------

    def _find_extent_for_blip(self, blip_element):
        """
        Remonte depuis un <a:blip> jusqu'à son conteneur <wp:inline>
        (image en ligne) ou <wp:anchor> (image flottante, ex :
        'devant le texte') pour en lire la taille d'origine
        <wp:extent cx=".." cy="..">, en EMU.
        """

        el = blip_element

        while el is not None:

            tag = el.tag

            if tag in (qn("wp:inline"), qn("wp:anchor")):

                extent = el.find(qn("wp:extent"))

                if extent is not None:
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    if cx and cy:
                        return int(cx), int(cy)

                return None

            el = el.getparent()

        return None

    # ------------------------------------------------------

    def _copy_paragraph_with_formatting(self, doc, src_paragraph):
        """
        Recopie un paragraphe source dans le rapport en conservant,
        run par run, le gras, l'italique, le souligné, la couleur,
        la police/taille, ainsi que les sauts de ligne manuels
        (Maj+Entrée) et tabulations à l'intérieur des runs — ET,
        au niveau du paragraphe lui-même, l'alignement, le retrait
        (gauche/droite/première ligne) et les taquets de
        tabulation définis (indispensable pour les paragraphes
        "en retrait", ex : listes d'espèces indentées).
        """

        new_p = doc.add_paragraph()
        new_p.alignment = src_paragraph.alignment

        src_pf = src_paragraph.paragraph_format
        dst_pf = new_p.paragraph_format

        dst_pf.left_indent = src_pf.left_indent
        dst_pf.right_indent = src_pf.right_indent
        dst_pf.first_line_indent = src_pf.first_line_indent
        dst_pf.space_before = src_pf.space_before
        dst_pf.space_after = src_pf.space_after
        dst_pf.line_spacing = src_pf.line_spacing

        try:
            for ts in src_pf.tab_stops:
                dst_pf.tab_stops.add_tab_stop(
                    ts.position, ts.alignment, ts.leader
                )
        except Exception:
            pass

        for run in src_paragraph.runs:

            new_run = new_p.add_run()

            # --- Mise en forme ---
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

            if run.font.size:
                new_run.font.size = run.font.size

            if run.font.name:
                new_run.font.name = run.font.name

            try:
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            except Exception:
                pass

            try:
                if run.font.highlight_color:
                    new_run.font.highlight_color = (
                        run.font.highlight_color
                    )
            except Exception:
                pass

            # --- Contenu : texte, sauts de ligne, tabulations,
            # dans l'ordre où ils apparaissent dans le run source ---
            for child in run._r:

                tag = child.tag

                if tag == qn("w:t"):
                    new_run.add_text(child.text or "")
                elif tag == qn("w:br"):
                    new_run.add_break()
                elif tag == qn("w:tab"):
                    new_run.add_tab()

        return new_p

    # ------------------------------------------------------

    def _copy_table_with_formatting(self, doc, src_table):
        """
        Copie un tableau natif d'un document source (ex : grille de
        sensibilité SFEPM, tableau des espèces) dans le rapport, en
        conservant les fusions de cellules (gridSpan horizontal),
        la couleur de fond de chaque cellule et la mise en forme du
        texte (gras/italique/souligné/police/couleur), run par run.
        """

        n_rows = len(src_table.rows)
        n_cols = len(src_table.columns)

        dst_table = doc.add_table(rows=n_rows, cols=n_cols)
        dst_table.style = "Table Grid"

        for i, src_row in enumerate(src_table.rows):

            # Cellules pointant vers le même <w:tc> source = fusion
            # horizontale à reproduire côté destination.
            tc_ids = [id(c._tc) for c in src_row.cells]

            j = 0
            while j < n_cols:

                src_cell = src_row.cells[j]

                span_end = j
                while (
                    span_end + 1 < n_cols
                    and tc_ids[span_end + 1] == tc_ids[j]
                ):
                    span_end += 1

                dst_cell = dst_table.cell(i, j)

                if span_end > j:
                    dst_cell = dst_cell.merge(
                        dst_table.cell(i, span_end)
                    )

                # --- Texte, paragraphe par paragraphe, run par run ---
                dst_cell.text = ""
                first = True
                for src_p in src_cell.paragraphs:
                    if first:
                        dst_p = dst_cell.paragraphs[0]
                        first = False
                    else:
                        dst_p = dst_cell.add_paragraph()
                    dst_p.alignment = src_p.alignment
                    for run in src_p.runs:
                        new_run = dst_p.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                        try:
                            if run.font.color and run.font.color.rgb:
                                new_run.font.color.rgb = (
                                    run.font.color.rgb
                                )
                        except Exception:
                            pass

                # --- Couleur de fond (shading) ---
                tcPr = src_cell._tc.tcPr
                fill = None
                if tcPr is not None:
                    shd = tcPr.find(qn("w:shd"))
                    if shd is not None:
                        fill = shd.get(qn("w:fill"))

                if fill and fill.lower() not in ("auto", "ffffff"):
                    dst_tcPr = dst_cell._tc.get_or_add_tcPr()
                    new_shd = OxmlElement("w:shd")
                    new_shd.set(qn("w:val"), "clear")
                    new_shd.set(qn("w:color"), "auto")
                    new_shd.set(qn("w:fill"), fill)
                    dst_tcPr.append(new_shd)

                j = span_end + 1

    # ------------------------------------------------------

    def _detect_existing_caption_numbers(self, src_doc):
        """
        Scanne le texte d'un document source (paragraphes ET
        cellules de tableaux) à la recherche de légendes numérotées
        écrites en dur par l'utilisateur, du type "Figure 3" ou
        "Tableau 2" (texte simple, pas un champ Word — celles-ci ne
        "parlent" pas au compteur automatique d'Albat). Retourne le
        plus grand numéro trouvé pour chaque label, ou 0 si aucun.
        """

        pattern = re.compile(
            r'\b(Figure|Tableau)\s+(\d+)\b', re.IGNORECASE
        )
        max_nums = {"Figure": 0, "Tableau": 0}

        def scan_text(text):
            for m in pattern.finditer(text):
                label = (
                    "Figure" if m.group(1).lower() == "figure"
                    else "Tableau"
                )
                num = int(m.group(2))
                if num > max_nums[label]:
                    max_nums[label] = num

        for p in src_doc.paragraphs:
            scan_text(p.text)

        for t in src_doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    scan_text(cell.text)

        return max_nums

    # ------------------------------------------------------

    def _seed_caption_counter(self, doc, label, value):
        """
        Insère un champ SEQ caché (\\r = réinitialise le compteur à
        la valeur donnée, \\h = n'affiche rien et n'apparaît pas
        dans la Liste des figures/tableaux) : les prochaines
        légendes numérotées automatiquement par Albat repartent à
        la suite de ce qui existe déjà en dur dans un fichier fourni
        par l'utilisateur (ex : "Figure 1" tapée à la main dans
        Statuts.docx -> les figures Albat suivantes démarrent à 2).
        """

        p = doc.add_paragraph()
        run = p.add_run()
        run.font.hidden = True
        r = run._r

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" SEQ {label} \\r {value} \\h "

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        r.append(fld_begin)
        r.append(instr)
        r.append(fld_end)

    # ------------------------------------------------------

    def _insert_materiel_description(self, doc, file_path):
        """
        Insère la description du modèle d'enregistreur choisi,
        lue depuis un fichier .txt (texte brut, un paragraphe par
        ligne non vide) ou .docx (texte ET images intégrées,
        paragraphe par paragraphe, dans l'ordre du fichier source,
        en conservant la taille d'origine des images (y compris les
        images flottantes / "devant le texte") et l'alignement
        d'origine des paragraphes — y compris justifié. Plusieurs
        images d'un même paragraphe source (ex : mises côte à côte
        via un positionnement flottant) sont réinsérées côte à côte
        dans un même paragraphe plutôt qu'empilées.
        """

        try:

            if file_path.lower().endswith(".docx"):

                src = Document(file_path)
                rIds_inserted = set()

                def _get_image_bytes(rId):
                    try:
                        return src.part.related_parts[rId].blob
                    except KeyError:
                        return None

                def _add_picture_with_fallback(run_img, image_bytes, pic_kwargs):
                    """Essaie l'insertion directe, puis une
                    conversion PNG (utile pour les images collées en
                    EMF/WMF). Retourne True si réussi."""

                    try:
                        run_img.add_picture(
                            BytesIO(image_bytes), **pic_kwargs
                        )
                        return True
                    except Exception:
                        pass

                    if PIL_OK:
                        try:
                            pil_img = PILImage.open(
                                BytesIO(image_bytes)
                            )
                            pil_img = pil_img.convert("RGB")
                            buf = BytesIO()
                            pil_img.save(buf, format="PNG")
                            buf.seek(0)
                            run_img.add_picture(buf, **pic_kwargs)
                            return True
                        except Exception:
                            pass

                    return False

                def _insert_images_group(rIds, alignment):
                    """Insère un groupe d'images (issues d'un même
                    paragraphe source) côte à côte, dans un seul
                    paragraphe de destination."""

                    rIds = [
                        r for r in rIds
                        if r and r not in rIds_inserted
                    ]

                    if not rIds:
                        return

                    p_img = doc.add_paragraph()
                    # Alignement d'origine du paragraphe source s'il
                    # est explicitement défini ; sinon on retombe
                    # sur la gauche (comportement par défaut réel de
                    # Word quand aucun alignement n'est précisé —
                    # PAS le centre, contrairement à ce qui était
                    # supposé ici avant correction).
                    p_img.alignment = (
                        alignment
                        if alignment is not None
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )

                    any_failed = False

                    for rId, blip_el in rIds:

                        image_bytes = _get_image_bytes(rId)

                        if image_bytes is None:
                            continue

                        extent = self._find_extent_for_blip(blip_el)

                        pic_kwargs = (
                            {
                                "width": Emu(extent[0]),
                                "height": Emu(extent[1])
                            }
                            if extent else {"width": Cm(10)}
                        )

                        run_img = p_img.add_run("  ")

                        ok = _add_picture_with_fallback(
                            run_img, image_bytes, pic_kwargs
                        )

                        if ok:
                            rIds_inserted.add(rId)
                        else:
                            any_failed = True

                    if not p_img.runs:
                        p_img._element.getparent().remove(p_img._element)

                    if any_failed:
                        self._add_table_placeholder(
                            doc,
                            "Une image du fichier de description n'a "
                            "pas pu être insérée (format non "
                            "supporté, fréquent pour une image "
                            "collée via Ctrl+V — enregistrée en "
                            "EMF/WMF). Essayez de la réinsérer dans "
                            "le fichier source via Insertion > Image "
                            "plutôt qu'un copier-coller."
                        )

                VML_NS = "urn:schemas-microsoft-com:vml"

                for block in _iter_block_items(src):

                    if isinstance(block, Table):
                        # Tableau natif du document source (ex :
                        # grille SFEPM) : copié tel quel, avec ses
                        # fusions et couleurs de fond.
                        self._copy_table_with_formatting(doc, block)
                        continue

                    p = block

                    # Copie le paragraphe même s'il est vide (une
                    # ligne blanche entre deux blocs de texte dans
                    # le fichier source est un choix de mise en
                    # forme délibéré : la sauter aplatit l'espacement
                    # d'origine). _copy_paragraph_with_formatting
                    # gère très bien un paragraphe sans runs.
                    self._copy_paragraph_with_formatting(doc, p)

                    # Cherche toutes les images de ce paragraphe (en
                    # ligne ou flottantes/"devant le texte"), format
                    # moderne (DrawingML) et ancien (VML), pour les
                    # regrouper et les insérer côte à côte.
                    group = []

                    for blip in p._p.findall(".//" + qn("a:blip")):
                        rId = (
                            blip.get(qn("r:embed"))
                            or blip.get(qn("r:link"))
                        )
                        if rId:
                            group.append((rId, blip))

                    for imgdata in p._p.findall(
                        f".//{{{VML_NS}}}imagedata"
                    ):
                        rId = imgdata.get(qn("r:id"))
                        if rId:
                            group.append((rId, imgdata))

                    _insert_images_group(group, p.alignment)

                # Filet de sécurité : images non rattachées à un
                # paragraphe détecté ci-dessus (mise en page
                # particulière) — ajoutées à la fin plutôt que
                # perdues silencieusement.
                remaining = []

                for shape in src.inline_shapes:
                    try:
                        rId = (
                            shape._inline.graphic.graphicData
                            .pic.blipFill.blip.embed
                        )
                        if rId not in rIds_inserted:
                            remaining.append(
                                (rId, shape._inline.graphic)
                            )
                    except Exception:
                        pass

                if remaining:
                    _insert_images_group(remaining, None)

                # NOTE : le recalage automatique du compteur
                # (détection de légendes en dur type "Figure 1"
                # dans le texte source) a été désactivé — trop de
                # faux positifs (une simple mention de "Figure 1"
                # dans une phrase, sans rapport avec une image
                # réellement insérée, décalait à tort la
                # numérotation de toutes les figures suivantes).
                # Les fonctions _detect_existing_caption_numbers et
                # _seed_caption_counter restent disponibles si
                # besoin, mais ne sont plus appelées automatiquement
                # ici — le workflow image+légende séparée n'en a
                # plus besoin.

            else:

                with open(file_path, "r", encoding="utf-8") as f:
                    contenu = f.read()

                for ligne in contenu.split("\n"):
                    if ligne.strip():
                        doc.add_paragraph(ligne.strip())

        except Exception as e:
            self._add_table_placeholder(
                doc,
                f"Impossible de lire le fichier de description "
                f"({os.path.basename(file_path)}) : {e}"
            )

    # ------------------------------------------------------

    def _add_table_placeholder(self, doc, caption):
        """Ajoute une légende de tableau/figure à insérer manuellement
        (données à coller depuis Correlations / Graph / Bridage)."""

        p = doc.add_paragraph()
        run = p.add_run(f"[TABLEAU/GRAPHIQUE À INSÉRER — {caption}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.size = Pt(10)
        return p

    # ------------------------------------------------------

    def _find_image(self, keyword, prefer_table=False):
        """
        Cherche, dans le dossier de graphiques indiqué, une image
        dont le nom contient le mot-clé donné (correspond aux noms
        générés par les boutons d'export de Graph, ex :
        'Parc_2025_03_Contacts_par_mois.jpg' ou
        'Parc_2025_T02_Contacts_par_mois.jpg').

        Certains exports existent en deux versions sous le même
        mot-clé (un graphique ET un tableau de données, ex :
        'Contacts_selon_vent'). Par défaut, priorité au graphique
        (plus visuel) ; prefer_table=True inverse cette priorité
        pour cibler explicitement la version tableau.

        La correspondance se fait sur la FIN du nom de fichier
        (avant l'extension), pas une sous-chaîne quelconque : sinon
        le mot-clé "Contacts_par_espece" matcherait aussi
        "Contacts_par_espece_et_mois.jpg" (qui le contient), un
        fichier totalement différent.
        """

        if not self.graphs_dir or not os.path.isdir(self.graphs_dir):
            return None

        try:
            files = os.listdir(self.graphs_dir)
        except Exception:
            return None

        keyword_low = keyword.lower()

        candidates = [
            f for f in files
            if f.lower().endswith((".jpg", ".jpeg", ".png"))
            and os.path.splitext(f)[0].lower().endswith(keyword_low)
        ]

        if not candidates:
            return None

        if prefer_table:
            preferred = [c for c in candidates if "_t0" in c.lower()]
        else:
            preferred = [c for c in candidates if "_t0" not in c.lower()]

        chosen = preferred[0] if preferred else candidates[0]

        return os.path.join(self.graphs_dir, chosen)

    # ------------------------------------------------------

    def _find_qgis_image(self):
        """
        Cherche, dans le dossier 'Qgis' du projet, l'image de
        localisation à insérer en I.3 (carte exportée depuis QGIS).
        Contrairement à _find_image (mot-clé requis, dossier
        Graph), ce dossier est dédié à cette seule carte : la
        première image trouvée (ordre alphabétique, fichiers
        cachés/verrous ignorés) est utilisée.
        """

        if not self.qgis_dir or not os.path.isdir(self.qgis_dir):
            return None

        try:
            files = sorted(os.listdir(self.qgis_dir))
        except Exception:
            return None

        for f in files:
            if f.startswith(".") or f.startswith("~$"):
                continue
            if f.lower().endswith((".jpg", ".jpeg", ".png")):
                return os.path.join(self.qgis_dir, f)

        return None

    # ------------------------------------------------------

    def _detect_materiel_utilise(self):
        """
        Déduit le nom du matériel à utiliser dans les phrases de la
        section I.3 ('Le Batcorder a été installé sur...') à partir
        du modèle d'enregistreur choisi dans le menu déroulant de
        la section Titre : détecte 'Batcorder' ou 'Batlogger' dans
        le nom du modèle choisi (insensible à la casse). Si aucun
        des deux n'est reconnu, réutilise tel quel le nom du modèle
        choisi ; à défaut, retombe sur un texte générique.
        """

        modele = self.combo_modele.currentText().strip()

        if not modele or modele.startswith("Aucun"):
            return "l'enregistreur"

        modele_low = modele.lower()

        if "batcorder" in modele_low:
            return "Batcorder"
        if "batlogger" in modele_low:
            return "Batlogger"

        return modele

    # ------------------------------------------------------

    def _find_scenar_file(self):
        """
        Cherche, dans le dossier 'Scenar' du projet, le fichier
        Excel le plus récemment exporté (plusieurs exports
        successifs sont possibles au fil des itérations d'un
        scénario ; on prend le plus récent). Ce fichier contient un
        onglet 'Synthèse' listant les plages du plan de bridage
        retenu (période, seuils vent/température, restriction
        nocturne, taux de protection).
        """

        if not self.scenar_dir or not os.path.isdir(self.scenar_dir):
            return None

        candidates = []

        try:
            for f in os.listdir(self.scenar_dir):
                if f.startswith("~$") or f.startswith("."):
                    continue
                if f.lower().endswith((".xlsx", ".xlsm")):
                    full = os.path.join(self.scenar_dir, f)
                    candidates.append((os.path.getmtime(full), full))
        except Exception:
            return None

        if not candidates:
            return None

        candidates.sort(reverse=True)

        return candidates[0][1]

    # ------------------------------------------------------

    def _insert_scenar_synthese(self, doc):
        """
        Lit l'onglet 'Synthèse' du dernier export Scenar du projet
        (plages du plan de bridage retenu, seuils, taux de
        protection) et l'insère sous forme de tableau Word natif,
        avec des largeurs de colonnes calculées d'après le contenu
        réel — pas de valeurs fixes — pour ne rien tronquer (la
        colonne 'Période' en particulier, plus longue que les
        autres colonnes).
        """

        if not OPENPYXL_OK:
            self._add_table_placeholder(
                doc,
                "Plan de régulation proposé — le module 'openpyxl' "
                "n'est pas installé, lecture du fichier Scenar "
                "impossible."
            )
            return

        scenar_file = self._find_scenar_file()

        if not scenar_file:
            self._add_table_placeholder(
                doc,
                "Plan de régulation proposé — aucun export Scenar "
                "trouvé dans le dossier '"
                f"{self.scenar_dir or 'Scenar (dossier du projet)'}"
                "'. Exportez le plan retenu depuis l'onglet Scenar."
            )
            return

        try:
            wb = openpyxl.load_workbook(scenar_file, data_only=True)
        except Exception as e:
            self._add_table_placeholder(
                doc,
                "Plan de régulation proposé — impossible de lire le "
                f"fichier Scenar ({os.path.basename(scenar_file)}) : "
                f"{e}"
            )
            return

        if "Synthèse" not in wb.sheetnames:
            self._add_table_placeholder(
                doc,
                "Plan de régulation proposé — onglet 'Synthèse' "
                f"introuvable dans '{os.path.basename(scenar_file)}'."
            )
            return

        ws = wb["Synthèse"]

        # Repère la ligne d'en-tête du tableau ("Plage", "Période"...)
        header_row_idx = None
        for r in range(1, ws.max_row + 1):
            if ws.cell(row=r, column=1).value == "Plage":
                header_row_idx = r
                break

        if header_row_idx is None:
            self._add_table_placeholder(
                doc,
                "Plan de régulation proposé — structure inattendue "
                "dans l'onglet 'Synthèse' du fichier Scenar."
            )
            return

        headers = [
            str(ws.cell(row=header_row_idx, column=c).value or "")
            for c in range(1, 7)
        ]

        data_rows = []
        r = header_row_idx + 1
        while True:
            first_cell = ws.cell(row=r, column=1).value
            if first_cell is None or str(first_cell).strip() == "":
                break
            data_rows.append([
                str(ws.cell(row=r, column=c).value or "")
                for c in range(1, 7)
            ])
            r += 1

        if not data_rows:
            self._add_table_placeholder(
                doc,
                "Plan de régulation proposé — aucune plage trouvée "
                "dans l'onglet 'Synthèse' du fichier Scenar."
            )
            return

        # Ligne "TAUX DE PROTECTION GLOBAL DU PLAN", quelques
        # lignes plus bas.
        taux_global = None
        for rr in range(r, min(r + 6, ws.max_row + 1)):
            val = ws.cell(row=rr, column=1).value
            if val and "TAUX DE PROTECTION GLOBAL" in str(val).upper():
                taux_global = ws.cell(row=rr, column=5).value
                break

        # --- Construction du tableau Word ---
        self._add_numbered_caption(
            doc, "Tableau", "Plan de régulation retenu"
        )

        n_cols = len(headers)

        max_lens = [
            max(
                [len(headers[j])]
                + [len(row[j]) for row in data_rows]
            )
            for j in range(n_cols)
        ]
        total_len = sum(max_lens) or 1

        # Largeur de page utile approximative (marges standard sur
        # une page A4 portrait).
        total_width_cm = 16.5
        col_widths = [
            max(total_width_cm * (l / total_len), 1.8)
            for l in max_lens
        ]
        s = sum(col_widths)
        col_widths = [w * total_width_cm / s for w in col_widths]

        table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
        table.style = "Table Grid"
        table.autofit = False

        def _style_cell(cell, text, width_cm, bold, fill_hex, font_color):
            cell.text = text
            cell.width = Cm(width_cm)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.bold = bold
                    if font_color:
                        run.font.color.rgb = font_color
            if fill_hex:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), fill_hex)
                tcPr.append(shd)

        for j, h in enumerate(headers):
            _style_cell(
                table.rows[0].cells[j], h, col_widths[j],
                bold=True, fill_hex="6E8F4E",
                font_color=RGBColor(0xFF, 0xFF, 0xFF)
            )

        for i, row_vals in enumerate(data_rows, start=1):
            for j, val in enumerate(row_vals):
                _style_cell(
                    table.rows[i].cells[j], val, col_widths[j],
                    bold=False, fill_hex=None, font_color=None
                )

        # python-docx exige de répéter la largeur sur chaque
        # cellule de chaque ligne ET sur table.columns pour que
        # Word (et LibreOffice) la respectent réellement (sinon un
        # partage égal est recalculé à l'ouverture).
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

        for j, col in enumerate(table.columns):
            col.width = Cm(col_widths[j])

        if taux_global:
            doc.add_paragraph()
            p_taux = doc.add_paragraph()
            run_taux = p_taux.add_run(
                f"Taux de protection global du plan : {taux_global}"
            )
            run_taux.bold = True

    # ------------------------------------------------------

    def _read_caption_text(self, docx_path):
        """
        Lit le texte d'un fichier Word ne contenant qu'une légende
        (un ou plusieurs paragraphes, concaténés). Fichier
        compagnon d'une image, même nom de base, ex :
        'schema.jpg' + 'schema.docx' (le texte de la légende).
        Retourne une chaîne vide si illisible.
        """

        try:
            d = Document(docx_path)
            lignes = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            return " ".join(lignes)
        except Exception:
            return ""

    # ------------------------------------------------------

    def _insert_numbered_image(
        self, doc, image_path, label="Figure",
        default_caption="", width_cm=14, max_height_cm=None
    ):
        """
        Insère une image centrée avec une légende Word numérotée
        automatiquement (champ SEQ, voir _add_numbered_caption). Le
        texte de la légende est cherché en priorité dans un fichier
        .docx compagnon (même nom de base que l'image, voir
        _resolve_caption) ; sinon default_caption. Mécanisme
        volontairement séparé image/texte — bien plus fiable
        qu'extraire une légende écrite en dur au milieu d'un
        fichier de description mixte.

        Pour les tableaux (label="Tableau"), une hauteur maximale
        par défaut est appliquée (voir plus bas) : certains
        tableaux sont étroits mais très hauts (peu de colonnes,
        beaucoup de lignes, ex : "Contacts selon vent"), et les
        insérer à pleine largeur (14 cm) les faisait prendre bien
        plus de hauteur que nécessaire, poussant leur légende sur
        la page suivante. On calcule la largeur réellement utilisée
        à partir des vraies proportions de l'image (largeur/hauteur
        en pixels), pour ne réduire QUE les images effectivement
        trop hautes — un tableau large mais peu haut (ex :
        "Indicateurs par mois", grille SFEPM) garde sa pleine
        largeur, donc sa lisibilité.
        """

        caption_text = self._resolve_caption(image_path, default_caption)

        if max_height_cm is None:
            max_height_cm = 10 if label == "Tableau" else None

        final_width_cm = width_cm

        if max_height_cm is not None:
            try:
                from PIL import Image as PILImage
                with PILImage.open(image_path) as im:
                    px_w, px_h = im.size
                if px_h > 0:
                    height_at_full_width = width_cm * (px_h / px_w)
                    if height_at_full_width > max_height_cm:
                        final_width_cm = max_height_cm * (px_w / px_h)
            except Exception:
                pass

        self._add_numbered_caption(doc, label, caption_text)

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_together = True
        run_img = p_img.add_run()

        try:
            run_img.add_picture(image_path, width=Cm(final_width_cm))
        except Exception:
            self._add_table_placeholder(doc, caption_text or "Image")
            return

    # ------------------------------------------------------

    def _find_influence_image(self, keyword):
        """
        Cherche, dans le dossier fixe 'images' (dossiers pour
        rapport/images), une image dont le nom contient le
        mot-clé donné (insensible à la casse). Contrairement aux
        graphiques de l'onglet Graph, ces figures sont des
        illustrations bibliographiques fixes, fournies par
        l'utilisateur (pas de problème de droits d'auteur puisque
        c'est son propre contenu) — pas liées au projet en cours.
        Ignore les fichiers cachés/verrous.
        """

        if not os.path.isdir(IMAGES_DIR):
            return None

        try:
            files = sorted(os.listdir(IMAGES_DIR))
        except Exception:
            return None

        for f in files:
            if f.startswith(".") or f.startswith("~$"):
                continue
            if not f.lower().endswith((".jpg", ".jpeg", ".png")):
                continue
            if keyword.lower() in f.lower():
                return os.path.join(IMAGES_DIR, f)

        return None

    # ------------------------------------------------------

    def _diagnose_images_dir(self):
        """
        Décrit l'état du dossier 'images' pour un message d'erreur
        exploitable : dossier absent, vide, ou liste des fichiers
        image qu'il contient réellement (pour repérer une extension
        non supportée ou un nom qui ne contient pas le mot-clé
        attendu, sans avoir à ouvrir l'explorateur de fichiers).
        """

        if not os.path.isdir(IMAGES_DIR):
            return "ce dossier n'existe pas encore."

        try:
            fichiers = [
                f for f in sorted(os.listdir(IMAGES_DIR))
                if not f.startswith(".") and not f.startswith("~$")
            ]
        except Exception:
            return "impossible de lire ce dossier."

        if not fichiers:
            return "ce dossier est vide."

        images = [
            f for f in fichiers
            if f.lower().endswith((".jpg", ".jpeg", ".png"))
        ]
        autres = [f for f in fichiers if f not in images]

        msg = f"fichiers présents dans ce dossier : {', '.join(fichiers)}."
        if autres:
            msg += (
                " Attention, seuls les .jpg/.jpeg/.png sont pris en "
                f"compte — ignoré(s) : {', '.join(autres)}."
            )
        return msg

    # ------------------------------------------------------

    def _insert_influence_image(self, doc, keyword, caption, label="Figure"):
        """
        Insère une image fixe unique (dossier images), centrée,
        avec légende numérotée — Figure ou Tableau selon le
        paramètre label (utile pour les tableaux fournis sous forme
        d'image, ex : grille de sensibilité SFEPM). Le texte de
        la légende vient en priorité d'un fichier .docx compagnon
        (même nom de base que l'image, ex : 'heure_nuit_1.jpg' +
        'heure_nuit_1.docx') ; à défaut, retombe sur le texte fourni
        en dur ici. Repli propre (repère à insérer manuellement) si
        l'image elle-même est introuvable.
        """

        image_path = self._find_influence_image(keyword)

        if not image_path:
            self._add_table_placeholder(
                doc,
                f"{caption} — aucun fichier ne contient '{keyword}' "
                f"dans le dossier '{IMAGES_DIR}' "
                f"({self._diagnose_images_dir()})"
            )
            return

        self._insert_numbered_image(
            doc, image_path, label=label, default_caption=caption
        )

    # ------------------------------------------------------

    def _insert_influence_images_side_by_side(
        self, doc, keyword_gauche, keyword_droite, caption
    ):
        """
        Insère deux figures bibliographiques fixes côte à côte
        (dossier images), via un tableau 1 ligne x 2 colonnes
        sans bordures, avec une légende commune en dessous. Si une
        seule des deux images est trouvée, elle est insérée seule
        (centrée) plutôt que de bloquer l'ensemble ; si aucune des
        deux n'est trouvée, un repère à insérer manuellement est
        laissé.
        """

        path_g = self._find_influence_image(keyword_gauche)
        path_d = self._find_influence_image(keyword_droite)

        if not path_g and not path_d:
            self._add_table_placeholder(
                doc,
                f"{caption} — aucun fichier ne contient "
                f"'{keyword_gauche}' ni '{keyword_droite}' dans le "
                f"dossier '{IMAGES_DIR}' "
                f"({self._diagnose_images_dir()})"
            )
            return

        if path_g and not path_d:
            self._insert_influence_image(doc, keyword_gauche, caption)
            return

        if path_d and not path_g:
            self._insert_influence_image(doc, keyword_droite, caption)
            return

        self._add_numbered_caption(
            doc, "Figure", self._resolve_caption(path_g, caption)
        )

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "nil")
                borders_el = tcPr.find(qn("w:tcBorders"))
                if borders_el is None:
                    borders_el = OxmlElement("w:tcBorders")
                    tcPr.append(borders_el)
                borders_el.append(border)

        for cell, img_path in zip(table.rows[0].cells, (path_g, path_d)):
            cell.width = Cm(8.2)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(img_path, width=Cm(7.6))
            except Exception:
                pass

    # ------------------------------------------------------

    def _resolve_caption(self, image_path, default_caption):
        """
        Cherche un fichier .docx compagnon (même nom de base que
        l'image) contenant le texte de légende à utiliser en
        priorité ; retombe sur default_caption si absent/vide.
        """

        if not image_path:
            return default_caption

        caption_path = os.path.splitext(image_path)[0] + ".docx"

        if os.path.exists(caption_path):
            texte = self._read_caption_text(caption_path)
            if texte:
                return texte

        return default_caption

    # ------------------------------------------------------

    def _insert_graphic(
        self, doc, keyword, caption, label="Figure", prefer_table=False,
        max_height_cm=None
    ):
        """
        Insère automatiquement l'image correspondant au mot-clé si
        elle est trouvée dans le dossier de graphiques choisi, avec
        légende numérotée (texte compagnon .docx prioritaire si
        présent, voir _resolve_caption), sinon laisse un repère à
        insérer manuellement. label="Tableau" pour les exports
        Graph qui sont en réalité des tableaux de données rendus en
        image (ex : "Indicateurs par mois"), pas des graphiques.
        prefer_table=True quand le même mot-clé existe en deux
        versions (graphique + tableau, ex : "Contacts_selon_vent")
        et qu'on veut explicitement la version tableau. max_height_cm
        permet d'augmenter la limite de hauteur par défaut (10 cm
        pour un tableau) quand celui-ci dispose de sa propre page
        entière (ex : après un saut de page dédié) et peut donc être
        affiché plus grand, plus lisible.
        """

        image_path = self._find_image(keyword, prefer_table=prefer_table)

        if image_path:
            self._insert_numbered_image(
                doc, image_path, label=label, default_caption=caption,
                max_height_cm=max_height_cm
            )
        else:
            self._add_table_placeholder(doc, caption)

    # ------------------------------------------------------

    def generate_report_template(self):

        if not DOCX_OK:
            QMessageBox.critical(
                self, "Module manquant",
                "Le module 'python-docx' n'est pas installé.\n"
                "Ouvrez un terminal et exécutez :\n"
                "pip install python-docx\n\n"
                "puis relancez l'application."
            )
            return

        modele_verif = self.combo_modele.currentText()

        if modele_verif.startswith("Aucun"):
            reply = QMessageBox.question(
                self, "Aucun matériel sélectionné",
                "Aucun modèle d'enregistreur n'a été sélectionné "
                "(section 'Titre de l'étude').\n\n"
                "La section I.1 'Matériel utilisé' contiendra un "
                "texte générique à compléter manuellement, sans la "
                "description ni la photo du matériel.\n\n"
                "Continuer quand même ?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply != QMessageBox.Yes:
                return

        try:

            bureau_nom = self.field_nom.text().strip() or "Bureau d'étude"
            adresse = self.field_adresse.toPlainText().strip()
            tel = self.field_tel.text().strip()
            email = self.field_email.text().strip()
            web = self.field_web.text().strip()

            titre_etude = self.field_titre_etude.text().strip()

            if not titre_etude:
                titre_etude = "Étude acoustique des chiroptères"

            parc_eolien = self.field_parc_eolien.text().strip()

            annee_txt = self.field_annee.text().strip()
            try:
                annee = int(annee_txt)
            except ValueError:
                annee = dtmod.date.today().year

            park_part = parc_eolien if parc_eolien else "Albat"

            safe_park = "".join(
                c if c.isalnum() or c in " _-" else "_"
                for c in park_part
            ).strip()

            default_name = f"{safe_park}_{annee}_Rapport.docx"

            if self.default_save_dir and os.path.isdir(self.default_save_dir):
                default_name = os.path.join(
                    self.default_save_dir, default_name
                )

            out_path, _ = QFileDialog.getSaveFileName(
                self,
                "Enregistrer le rapport",
                default_name,
                "Word (*.docx)"
            )

            if not out_path:
                return

            if not out_path.lower().endswith(".docx"):
                out_path += ".docx"

            doc = Document()

            # ==================================================
            # STYLES DE BASE
            # ==================================================

            corps_police = self.font_combo.currentFont().family() or "Calibri"

            taille_titre = self.spin_size_titre.value()
            taille_soustitre = self.spin_size_soustitre.value()
            taille_corps = self.spin_size_corps.value()

            couleur_titre_rgb = RGBColor(
                self.color_titre.red(),
                self.color_titre.green(),
                self.color_titre.blue()
            )
            couleur_corps_rgb = RGBColor(
                self.color_corps.red(),
                self.color_corps.green(),
                self.color_corps.blue()
            )

            normal = doc.styles["Normal"]
            normal.font.name = corps_police
            normal.font.size = Pt(taille_corps)
            normal.font.color.rgb = couleur_corps_rgb

            # Neutralise l'espacement par défaut de Word (souvent
            # ~8pt avant/après chaque paragraphe), qui s'accumule et
            # peut faire déborder la page de garde sur 2 pages.
            normal.paragraph_format.space_before = Pt(0)
            normal.paragraph_format.space_after = Pt(0)

            # S'assure que la police s'applique aussi côté polices
            # non-occidentales (nécessaire pour un rendu cohérent
            # dans Word, notamment sur Mac).
            rpr = normal.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), corps_police)

            heading_sizes = {
                "Heading 1": taille_titre,
                "Heading 2": taille_soustitre,
                "Heading 3": max(taille_soustitre - 1, 6),
            }

            for heading_name, size in heading_sizes.items():
                try:
                    doc.styles[heading_name].font.name = corps_police
                    doc.styles[heading_name].font.size = Pt(size)
                    doc.styles[heading_name].font.color.rgb = (
                        couleur_titre_rgb
                    )
                except KeyError:
                    pass

            # ==================================================
            # PAGE DE GARDE
            # ==================================================

            mois_fr = {
                1: "Janvier", 2: "Février", 3: "Mars", 4: "Avril",
                5: "Mai", 6: "Juin", 7: "Juillet", 8: "Août",
                9: "Septembre", 10: "Octobre", 11: "Novembre",
                12: "Décembre"
            }
            mois_actuel = mois_fr[dtmod.date.today().month]

            # --- Logo principal ---
            if self.logo_path and os.path.exists(self.logo_path):
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.paragraph_format.space_before = Pt(18)
                p_logo.paragraph_format.space_after = Pt(24)
                run_logo = p_logo.add_run()
                try:
                    run_logo.add_picture(self.logo_path, width=Cm(8))
                except Exception:
                    pass
            else:
                p_bureau = doc.add_paragraph()
                p_bureau.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_bureau.paragraph_format.space_before = Pt(18)
                p_bureau.paragraph_format.space_after = Pt(24)
                run_bureau = p_bureau.add_run(bureau_nom.upper())
                run_bureau.bold = True
                run_bureau.font.size = Pt(22)
                run_bureau.font.name = corps_police

            # --- Séparateur ---
            p_sep = doc.add_paragraph()
            p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sep.paragraph_format.space_after = Pt(40)
            run_sep = p_sep.add_run("–")
            run_sep.font.size = Pt(14)

            # --- Titre de l'étude (majuscules, plusieurs lignes) ---
            p_titre = doc.add_paragraph()
            p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_titre.paragraph_format.space_after = Pt(10)
            run_titre = p_titre.add_run(titre_etude.upper())
            run_titre.bold = False
            run_titre.font.size = Pt(22)
            run_titre.font.name = corps_police
            run_titre.font.color.rgb = couleur_titre_rgb

            # --- Parc éolien (sous le titre) ---
            if parc_eolien:
                p_parc = doc.add_paragraph()
                p_parc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_parc.paragraph_format.space_after = Pt(20)
                run_parc = p_parc.add_run(f"Parc de {parc_eolien}")
                run_parc.font.size = Pt(16)
                run_parc.font.name = corps_police
                run_parc.font.color.rgb = couleur_titre_rgb

            # --- "Rapport d'expertise, Mois Année" ---
            p_cr = doc.add_paragraph()
            p_cr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cr.paragraph_format.space_after = Pt(50)
            run_cr = p_cr.add_run(
                f"Rapport d'expertise, {mois_actuel} {annee}"
            )
            run_cr.font.size = Pt(18)
            run_cr.italic = True
            run_cr.font.name = corps_police

            # --- Petit logo décoratif en bas (logo secondaire si
            # défini, sinon le logo principal en miniature) ---
            petit_logo = (
                self.logo_secondaire_path
                if self.logo_secondaire_path
                and os.path.exists(self.logo_secondaire_path)
                else self.logo_path
            )

            if petit_logo and os.path.exists(petit_logo):
                p_logo2 = doc.add_paragraph()
                p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo2.paragraph_format.space_after = Pt(6)
                run_logo2 = p_logo2.add_run()
                try:
                    run_logo2.add_picture(petit_logo, width=Cm(1.8))
                except Exception:
                    pass

            # --- Coordonnées (avec libellés) ---
            coord_lines = [
                ("Adresse : ", adresse),
                ("Tel : ", tel),
                ("Mail : ", email),
                ("Site internet : ", web),
            ]

            for label, value in coord_lines:
                if value:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f"{label}{value}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            doc.add_page_break()

            # ==================================================
            # SOMMAIRE
            # ==================================================

            h_som = doc.add_heading("SOMMAIRE", level=1)

            self._add_toc_field(doc)

            doc.add_page_break()

            # ==================================================
            # LISTE DES FIGURES / LISTE DES TABLEAUX
            # ==================================================

            doc.add_heading("LISTE DES FIGURES", level=1)
            self._add_list_field(doc, "Figure")

            doc.add_page_break()

            doc.add_heading("LISTE DES TABLEAUX", level=1)
            self._add_list_field(doc, "Tableau")

            doc.add_page_break()

            # ==================================================
            # INTRODUCTION
            # ==================================================

            doc.add_heading("Introduction", level=1)

            annee_intro = self.field_annee.text().strip() or "[ANNÉE]"
            parc_intro = parc_eolien or "[NOM DU PARC]"

            p_intro = doc.add_paragraph()
            p_intro.add_run(
                TEXTES_SILVA["introduction_p1"].format(
                    parc=parc_intro, annee=annee_intro
                )
            )
            doc.add_paragraph(
                TEXTES_SILVA["introduction_p2"].format(
                    parc=parc_intro, annee=annee_intro
                )
            )

            todo_intro = doc.add_paragraph()
            run_todo_intro = todo_intro.add_run(
                TEXTES_SILVA["introduction_todo"]
            )
            run_todo_intro.italic = True
            run_todo_intro.bold = True
            run_todo_intro.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
            run_todo_intro.font.size = Pt(self.spin_size_corps.value())

            doc.add_page_break()

            # ==================================================
            # I. MATÉRIEL ET MÉTHODE
            # ==================================================

            doc.add_heading("I. Matériel et méthode", level=1)

            doc.add_paragraph(TEXTES_SILVA["materiel_methode_p1"])
            doc.add_paragraph(TEXTES_SILVA["materiel_methode_p2"])
            doc.add_paragraph(TEXTES_SILVA["materiel_methode_p3"])
            doc.add_paragraph(TEXTES_SILVA["materiel_methode_p4"])

            doc.add_heading("I.1. Matériel utilisé", level=2)
            doc.add_paragraph(TEXTES_SILVA["materiel_utilise_intro"])

            modele_choisi = self.combo_modele.currentText()

            # 1) Texte descriptif du matériel (fichier
            # '<Modèle>.docx') — sans l'éventuelle image, qui vient
            # séparément juste après (voir 3).
            if (
                modele_choisi in self.materiel_files
                and os.path.exists(self.materiel_files[modele_choisi])
            ):
                self._insert_materiel_description(
                    doc, self.materiel_files[modele_choisi]
                )
            else:
                self._add_todo(
                    doc,
                    "précisez ici le modèle d'enregistreur utilisé et "
                    "ses spécificités techniques le cas échéant."
                )

            # 2) Phrase de période de fonctionnement, réutilisant
            # les champs déjà saisis en I.3 (Période d'étude et
            # éolienne équipée) — pas de ressaisie. Positionnée
            # entre le texte descriptif et l'image (voir 3), comme
            # dans le rapport type. Si votre fichier "matériel
            # utilisé" contient déjà une phrase similaire du type
            # "L'enregistreur a fonctionné en continu du [à
            # compléter]...", supprimez-la du fichier Word source :
            # elle est désormais générée automatiquement ici.
            date_install_i1 = self.field_date_install.text().strip()
            date_desinstall_i1 = self.field_date_desinstall.text().strip()
            eoliennes_i1 = self.field_eoliennes.text().strip()

            p_fonctionnement = doc.add_paragraph()
            p_fonctionnement.add_run("L'enregistreur a fonctionné en continu du ")

            if date_install_i1:
                p_fonctionnement.add_run(date_install_i1)
            else:
                run_todo1 = p_fonctionnement.add_run("[à compléter]")
                run_todo1.bold = True
                run_todo1.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

            p_fonctionnement.add_run(" au ")

            if date_desinstall_i1:
                p_fonctionnement.add_run(date_desinstall_i1)
            else:
                run_todo2 = p_fonctionnement.add_run("[à compléter]")
                run_todo2.bold = True
                run_todo2.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

            p_fonctionnement.add_run(", sur l'éolienne ")

            if eoliennes_i1:
                p_fonctionnement.add_run(eoliennes_i1)
            else:
                run_todo3 = p_fonctionnement.add_run("[à compléter]")
                run_todo3.bold = True
                run_todo3.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

            p_fonctionnement.add_run(".")

            # 3) Image + légende du matériel (fichier compagnon
            # '<Modèle>_figure.jpg/.png', ou ancien format
            # '<Modèle>_figure.docx' contenant image+texte),
            # optionnel. Absent -> rien n'est ajouté (pas d'erreur,
            # ce n'est pas obligatoire).
            figure_file = self.materiel_figure_files.get(
                modele_choisi
            ) or self.materiel_figure_files.get("__default__")

            if figure_file and os.path.exists(figure_file):
                if figure_file.lower().endswith((".jpg", ".jpeg", ".png")):
                    self._insert_numbered_image(
                        doc, figure_file, label="Figure",
                        default_caption=(
                            f"Matériel utilisé — {modele_choisi}"
                        )
                    )
                else:
                    # Ancien format : image et texte mêlés dans le
                    # même .docx, inséré tel quel (pas de légende
                    # numérotée automatique dans ce cas précis).
                    self._insert_materiel_description(doc, figure_file)

            doc.add_heading("I.2. Analyse des enregistrements", level=2)

            analyse_file = self._find_analyse_file()

            if analyse_file and os.path.exists(analyse_file):
                self._insert_materiel_description(doc, analyse_file)
            else:
                self._add_table_placeholder(
                    doc,
                    "Analyse des enregistrements — fichier "
                    "introuvable. Déposez un fichier .docx dans le "
                    f"dossier intégré '{ANALYSE_DIR}'."
                )
                self._add_todo(
                    doc,
                    "décrivez le logiciel utilisé pour l'analyse des "
                    "enregistrements et la méthode de quantification "
                    "de l'activité (nombre de contacts)."
                )

            doc.add_heading(
                "I.3. Période d'étude et éolienne équipée", level=2
            )

            eoliennes = (
                self.field_eoliennes.text().strip()
                or "[éolienne(s) équipée(s)]"
            )
            date_install = (
                self.field_date_install.text().strip()
                or "[date d'installation]"
            )
            date_desinstall = (
                self.field_date_desinstall.text().strip()
                or "[date de désinstallation]"
            )
            annee_comparaison = (
                self.field_annee_suivi_anterieur.text().strip()
            )
            materiel_nom = self._detect_materiel_utilise()
            parc_nom = parc_eolien or "[nom du parc]"

            doc.add_paragraph(
                f"Le dispositif a été installé sur {eoliennes} au "
                "niveau du plancher de la nacelle. Ceci est "
                "conforme au protocole de suivi environnemental "
                "des parcs éoliens terrestres (révision 2018) qui "
                "impose un enregistreur pour 8 éoliennes. La pose "
                f"du matériel est intervenue le {date_install}. La "
                "désinstallation du matériel est intervenue le "
                f"{date_desinstall}."
            )

            doc.add_paragraph(
                f"Le {materiel_nom} a été installé sur {eoliennes} "
                "pour trois raisons :"
            )

            raison1 = doc.add_paragraph(style="List Bullet")
            raison1.add_run(
                "Parce que les milieux présents à proximité de "
                f"{eoliennes} sont représentatifs des milieux "
                f"observés sur l'ensemble du parc éolien "
                f"{parc_nom}"
            )

            raison2 = doc.add_paragraph(style="List Bullet")
            raison2.add_run(
                "Pour sa position centrale au sein du parc éolien "
                f"{parc_nom}"
            )

            raison3 = doc.add_paragraph(style="List Bullet")
            if annee_comparaison:
                raison3.add_run(
                    "Par soucis de comparaison avec le suivi "
                    f"réalisé en {annee_comparaison}"
                )
            else:
                raison3.add_run(
                    "Par soucis de comparaison avec le suivi "
                    "réalisé en "
                )
                run_todo = raison3.add_run("[À COMPLÉTER]")
                run_todo.bold = True
                run_todo.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

            qgis_image = self._find_qgis_image()
            caption_qgis = (
                f"Localisation du {materiel_nom} sur le parc de "
                f"{parc_nom}"
            )

            if qgis_image:
                self._insert_numbered_image(
                    doc, qgis_image, label="Figure",
                    default_caption=caption_qgis
                )
            else:
                self._add_table_placeholder(
                    doc,
                    f"{caption_qgis} — image introuvable, déposez "
                    "la carte exportée depuis QGIS dans le dossier "
                    f"'{self.qgis_dir or 'Qgis (dossier du projet)'}'."
                )

            doc.add_page_break()

            # ==================================================
            # II. RESULTATS
            # ==================================================

            doc.add_heading("II. Résultats", level=1)

            doc.add_heading("II.1. Activité enregistrée", level=2)

            resume = self._load_resume_rapport()

            if resume and resume.get("total_contacts") is not None:

                eoliennes_txt = (
                    self.field_eoliennes.text().strip()
                    or "[éolienne(s) équipée(s)]"
                )

                # Nombre de nuits d'écoute (dénominateur) : on
                # préfère les dates d'installation/désinstallation
                # saisies en I.3 (période réelle de suivi) ; à
                # défaut, on retombe sur l'écart entre le premier
                # et le dernier contact enregistré (JSON de Graph).
                nb_nuits_ecoute = None

                d_install = self._parse_date_fr(
                    self.field_date_install.text().strip()
                )
                d_desinstall = self._parse_date_fr(
                    self.field_date_desinstall.text().strip()
                )

                if d_install and d_desinstall and d_desinstall >= d_install:
                    # Pas de "+1" : le nombre de nuits entre deux
                    # dates de calendrier est l'écart en jours, pas
                    # l'écart+1 (vérifié sur le rapport de
                    # référence Silva : 26/02 -> 06/11 = 253 nuits,
                    # exactement l'écart en jours).
                    nb_nuits_ecoute = (d_desinstall - d_install).days
                elif resume.get("nb_jours_couverts_estime"):
                    nb_nuits_ecoute = resume["nb_jours_couverts_estime"]

                total_contacts = resume["total_contacts"]
                nb_mois = resume.get("nb_mois_couverts")
                nb_positives = resume.get("nb_nuits_positives")

                p_activite = doc.add_paragraph()
                p_activite.add_run(
                    f"{total_contacts} contacts ont été enregistrés "
                    f"sur {eoliennes_txt}"
                )
                if nb_mois:
                    p_activite.add_run(f" en {nb_mois} mois d'étude")
                p_activite.add_run(
                    ". Le nombre de nuits où nous avons obtenu des "
                    "enregistrements de chauves-souris est qualifié "
                    "de "
                )
                run_qualif = p_activite.add_run(
                    "[faible / modéré / important — à choisir]"
                )
                run_qualif.bold = True
                run_qualif.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
                if nb_positives is not None and nb_nuits_ecoute:
                    p_activite.add_run(
                        f" ({nb_positives} sur {nb_nuits_ecoute} "
                        "nuits d'écoute)."
                    )
                else:
                    p_activite.add_run(".")

            else:
                self._add_todo(
                    doc,
                    "nombre total de contacts enregistrés, nombre de "
                    "nuits positives sur nuits d'écoute, niveau "
                    "d'activité global (exportez les tableaux depuis "
                    "l'onglet Graph pour un remplissage automatique)."
                )

            self._insert_graphic(
                doc,
                "Contacts_par_mois",
                "Nombre de contacts par mois"
            )
            self._insert_graphic(
                doc,
                "Indicateurs",
                "Nombre de contacts et de nuits positives par mois",
                label="Tableau"
            )
            self._add_todo(
                doc,
                "commentez la répartition de l'activité dans le temps "
                "(pics d'activité, périodes de plus forte activité)."
            )
            self._insert_graphic(
                doc,
                "Contacts_par_date",
                "Nombre de contacts en fonction du temps"
            )

            doc.add_heading("II.2. Diversité spécifique", level=2)

            especes_data = resume.get("especes") if resume else None

            if especes_data:

                especes_triees = sorted(
                    especes_data.items(),
                    key=lambda kv: kv[1],
                    reverse=True
                )

                total_contacts_especes = (
                    resume.get("total_contacts")
                    or sum(c for _, c in especes_triees)
                )

                def _fmt_pct(valeur, total):
                    return str(
                        round(100 * valeur / total, 1)
                    ).replace(".", ",")

                doc.add_paragraph(
                    f"Au moins {len(especes_triees)} espèces ont "
                    "été identifiées lors du suivi mené en nacelle :"
                )

                for nom_espece, _ in especes_triees:
                    p_esp = doc.add_paragraph(style="List Bullet")
                    p_esp.add_run(nom_espece)

                doc.add_paragraph(
                    "Sur l'ensemble de la période d'étude, il a "
                    f"été comptabilisé {total_contacts_especes} "
                    "contacts de chiroptères."
                )

                if especes_triees and total_contacts_especes:
                    nom_top, contacts_top = especes_triees[0]
                    p_majoritaire = doc.add_paragraph()
                    p_majoritaire.add_run(
                        f"L'espèce la plus représentée est {nom_top}, "
                        "avec "
                        f"{_fmt_pct(contacts_top, total_contacts_especes)}"
                        "% des contacts totaux"
                        f" ({contacts_top} contacts)."
                    )

                    if len(especes_triees) > 1:
                        nom_2e, contacts_2e = especes_triees[1]
                        p_2e = doc.add_paragraph()
                        p_2e.add_run(
                            "La deuxième espèce la plus représentée "
                            f"est {nom_2e}, avec "
                            f"{_fmt_pct(contacts_2e, total_contacts_especes)}"
                            f"% des contacts totaux ({contacts_2e} "
                            "contacts)."
                        )

            else:
                self._add_todo(
                    doc,
                    "listez les espèces identifiées et leur nombre "
                    "de contacts respectif (exportez les tableaux "
                    "depuis l'onglet Graph pour un remplissage "
                    "automatique)."
                )

            self._insert_graphic(
                doc,
                "Contacts_par_espece",
                "Nombre de contacts par espèce recensée"
            )
            self._insert_graphic(
                doc,
                "Contacts_par_espece_et_mois",
                "Répartition des contacts par espèce et par mois"
            )
            self._insert_graphic(
                doc,
                "Activite_taxons",
                "Activité comparée par groupe d'espèces"
            )
            self._add_todo(
                doc,
                "commentez les regroupements d'espèces en cas "
                "d'identification incertaine (ex : groupe "
                "Noctules/Sérotines) et la répartition de "
                "l'activité par espèce au fil des mois de suivi."
            )

            doc.add_heading(
                "II.3. Statuts des espèces et sensibilité à l'éolien",
                level=2
            )

            # Découpage en 4 blocs de texte (fichiers "part1" à
            # "part4" dans le dossier Statuts des espèces),
            # entrecoupés de deux tableaux et d'une figure fournis
            # en image dans le dossier "images". L'ancien fichier
            # unique "Statuts.docx" n'est volontairement plus
            # utilisé ici (conservé tel quel dans son dossier en
            # attendant de valider que cette nouvelle structure
            # fonctionne bien).
            def _insert_statuts_part(numero):
                part_file = self._find_statuts_part_file(
                    f"part{numero}"
                )
                if part_file and os.path.exists(part_file):
                    self._insert_materiel_description(doc, part_file)
                else:
                    self._add_table_placeholder(
                        doc,
                        f"Statuts des espèces (partie {numero}) — "
                        f"fichier introuvable. Déposez, dans le "
                        f"dossier intégré '{STATUTS_DIR}', un "
                        f"fichier dont le nom contient "
                        f"'part{numero}'."
                    )

            _insert_statuts_part(1)

            self._insert_influence_image(
                doc, "niveau_sensibilite",
                "Définition des niveaux de sensibilité des espèces "
                "à l'éolien (SFEPM)",
                label="Tableau"
            )

            _insert_statuts_part(2)

            self._insert_influence_image(
                doc, "niveau_risque",
                "Évaluation du niveau de risque des différentes "
                "espèces identifiées en hauteur",
                label="Tableau"
            )

            _insert_statuts_part(3)

            self._insert_influence_image(
                doc, "routes_migratoires",
                "Routes migratoires",
                label="Figure"
            )

            _insert_statuts_part(4)

            doc.add_heading(
                "II.4. Corrélation de l'activité avec les paramètres "
                "biologiques et environnementaux",
                level=2
            )

            doc.add_heading("Influence de la température", level=3)
            doc.add_paragraph(TEXTES_SILVA["influence_temperature_p1"])
            doc.add_paragraph(TEXTES_SILVA["influence_temperature_p2"])
            self._add_todo(doc, "seuil de température observé sur le site.")
            self._insert_graphic(
                doc,
                "Contacts_selon_temperature",
                "Contacts par classe de température"
            )
            doc.add_page_break()
            self._insert_graphic(
                doc,
                "Contacts_selon_temperature",
                "Nombre de contacts en fonction de la température",
                label="Tableau",
                prefer_table=True,
                max_height_cm=19
            )

            doc.add_heading("Influence de la vitesse de vent", level=3)
            doc.add_paragraph(TEXTES_SILVA["influence_vent_p1"])
            self._add_todo(doc, "seuil de vent observé sur le site.")
            self._insert_graphic(
                doc,
                "Contacts_selon_vent",
                "Contacts par classe de vitesse de vent"
            )
            doc.add_page_break()
            self._insert_graphic(
                doc,
                "Contacts_selon_vent",
                "Nombre de contacts en fonction de la vitesse du vent",
                label="Tableau",
                prefer_table=True,
                max_height_cm=19
            )

            doc.add_heading("Influence de l'heure de la nuit", level=3)
            doc.add_paragraph(TEXTES_SILVA["influence_heure_nuit_p1"])

            self._insert_influence_image(
                doc,
                "heure_nuit_1",
                "Activité en fonction de l'heure de la nuit "
                "(synthèse bibliographique)"
            )

            doc.add_paragraph(TEXTES_SILVA["influence_heure_nuit_p2"])

            self._insert_influence_images_side_by_side(
                doc,
                "heure_nuit_2a",
                "heure_nuit_2b",
                "Distribution des contacts en fonction de l'heure de "
                "la nuit (synthèse bibliographique)"
            )

            self._add_todo(
                doc,
                "période de la nuit la plus active, part de l'activité "
                "avant le coucher / après le lever du soleil."
            )
            self._insert_graphic(
                doc,
                "Contacts_apres_coucher",
                "Contacts par heure après le coucher du soleil"
            )

            doc.add_page_break()

            # ==================================================
            # III. COMPARAISON AVEC UN SUIVI ANTÉRIEUR (optionnel)
            #
            # Insérée seulement si la case correspondante est
            # cochée dans la section "Comparaison avec un suivi
            # antérieur" de l'interface — sinon on passe
            # directement à la IV, comme dans le modèle type par
            # défaut.
            # ==================================================

            if self.check_comparaison.isChecked():

                annee_anterieure = (
                    self.field_annee_suivi_anterieur.text().strip()
                    or "[ANNÉE DU SUIVI ANTÉRIEUR]"
                )

                doc.add_heading(
                    "III. Comparaison avec le suivi "
                    f"{annee_anterieure}", level=1
                )

                self._insert_graphic(
                    doc,
                    "Comparaison_total",
                    "Nombre total de contacts enregistrés — "
                    f"comparaison {annee_anterieure} / suivi actuel"
                )
                self._add_todo(
                    doc,
                    "commentez l'écart de niveau d'activité entre "
                    "les deux années de suivi."
                )

                self._insert_graphic(
                    doc,
                    "Comparaison_evolution",
                    "Évolution de l'activité au fil des mois — "
                    f"comparaison {annee_anterieure} / suivi actuel"
                )
                self._add_todo(
                    doc,
                    "commentez les différences de répartition "
                    "saisonnière de l'activité entre les deux "
                    "années."
                )

                self._insert_graphic(
                    doc,
                    "Comparaison_groupes",
                    "Proportion des groupes d'espèces — comparaison "
                    f"{annee_anterieure} / suivi actuel"
                )

                self._insert_graphic(
                    doc,
                    "Comparaison_especes",
                    "Distribution des contacts par espèce — "
                    f"comparaison {annee_anterieure} / suivi actuel"
                )
                self._add_todo(
                    doc,
                    "commentez les évolutions par espèce (espèce "
                    "devenue majoritaire, espèces nouvellement "
                    "contactées, disparitions apparentes...) et les "
                    "explications possibles (changement de gîte, "
                    "effet vase communicant entre colonies "
                    "voisines...)."
                )

                doc.add_page_break()

            # ==================================================
            # IV. SYNTHESE ET DISCUSSION
            #
            # Numérotée IV même quand la III (comparaison) est
            # absente, pour rester cohérente avec le modèle type :
            # la numérotation ne se décale pas selon que la section
            # optionnelle est incluse ou non.
            # ==================================================

            doc.add_heading("IV. Synthèse et discussion", level=1)
            self._add_todo(
                doc,
                "bilan général de l'activité observée, explication des "
                "variations saisonnières, espèces les plus concernées."
            )

            doc.add_paragraph()
            p_mesure = doc.add_paragraph()
            run_mesure = p_mesure.add_run(
                "Plan de régulation proposé"
            )
            run_mesure.bold = True

            self._insert_scenar_synthese(doc)

            doc.add_page_break()

            # ==================================================
            # BIBLIOGRAPHIE
            # ==================================================

            doc.add_heading("Bibliographie", level=1)

            biblio_checked = [
                self.biblio_list.item(i).text()
                for i in range(self.biblio_list.count())
                if self.biblio_list.item(i).checkState() == Qt.Checked
            ]

            if biblio_checked:
                for ref in biblio_checked:
                    p_ref = doc.add_paragraph(ref)
                    p_ref.paragraph_format.space_after = Pt(12)
            else:
                self._add_todo(
                    doc,
                    "listez ici les références bibliographiques citées "
                    "dans le rapport (ou cochez-en dans le bloc "
                    "Bibliographie avant de générer)."
                )

            # Demande à Word de mettre à jour tous les champs
            # (sommaire, listes de figures/tableaux, numéros de
            # légende) dès l'ouverture du document, au lieu de
            # devoir faire un clic droit sur chacun séparément.
            # Word affichera une invite "Ce document contient des
            # champs... voulez-vous les mettre à jour ?" à l'ouverture.
            try:
                settings_el = doc.settings.element
                update_fields_el = OxmlElement("w:updateFields")
                update_fields_el.set(qn("w:val"), "true")
                settings_el.append(update_fields_el)
            except Exception:
                pass

            doc.save(out_path)

            open_file(out_path)

            QMessageBox.information(
                self, "Succès",
                f"Rapport type généré :\n{out_path}\n\n"
                f"Pensez à faire un clic droit sur le sommaire dans "
                f"Word puis « Mettre à jour les champs » pour qu'il "
                f"se remplisse automatiquement."
            )

        except Exception as e:
            QMessageBox.critical(
                self, "Erreur",
                f"Impossible de générer le rapport :\n{e}"
            )


if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = RapportWindow()
    win.show()
    sys.exit(app.exec())